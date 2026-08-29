# -*- coding: utf-8 -*-
from __future__ import annotations

import json
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT_DIR = Path(__file__).resolve().parent.parent
RULES_DIR = ROOT_DIR / "0_REGRAS_MANDATO"
BASE_JSON_DIR = RULES_DIR / "BASE_JSON"
OUTPUT_PATH = RULES_DIR / "PROMPT_GERAL_v3_ARCHANGEL_ROADMAP_INTERNO.docx"


def load_json(name: str) -> dict:
    path = BASE_JSON_DIR / name
    if not path.is_file():
        return {}
    try:
        with path.open("r", encoding="utf-8") as handle:
            data = json.load(handle)
        return data if isinstance(data, dict) else {}
    except Exception:
        return {}


def nested(data: dict, *keys, default=None):
    current = data
    for key in keys:
        if not isinstance(current, dict):
            return default
        current = current.get(key)
    return current if current is not None else default


def add_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths_in):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths_in):
            if idx >= len(row.cells):
                continue
            row.cells[idx].width = Inches(width)
            set_cell_margins(row.cells[idx])
            row.cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.TOP


def style_run(run, bold=False, italic=False, size=None, color=None):
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")


def add_para(doc, text="", style=None, bold_prefix=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.18
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        style_run(r, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        style_run(r2)
    else:
        r = p.add_run(text)
        style_run(r)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    style_run(run)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    style_run(run)
    return p


def add_h(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(5)
    for run in p.runs:
        style_run(run, bold=True, size=16 if level == 1 else 13, color="2E74B5" if level <= 2 else "1F4D78")
    return p


def add_status_table(doc, rows):
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, label in enumerate(["Camada", "Estado Atual", "Risco", "Próxima Ação"]):
        hdr[idx].text = label
        add_shading(hdr[idx], "E8EEF5")
        for run in hdr[idx].paragraphs[0].runs:
            style_run(run, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)
            for p in cells[idx].paragraphs:
                for run in p.runs:
                    style_run(run, size=9.5)
    set_table_widths(table, [1.45, 1.7, 1.35, 2.0])
    return table


def add_two_col_table(doc, rows):
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, detail in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = detail
        add_shading(cells[0], "F2F4F7")
        for run in cells[0].paragraphs[0].runs:
            style_run(run, bold=True)
        for cell in cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    style_run(run, size=9.5)
    set_table_widths(table, [1.7, 4.8])
    return table


def add_page_break(doc):
    doc.add_page_break()


def build_doc():
    ai_index = load_json("ARCHANGEL_AI_CONTEXT_INDEX.json")
    run_all = load_json("00_RUN_ALL_RUN_REPORT_LATEST.json")
    dq = load_json("DATA_QUALITY_REPORT.json")
    ds = load_json("5_JSON_DATASETS_ML.json")
    wf = load_json("6_JSON_WALK_FORWARD.json")
    pyenv = load_json("ARCHANGEL_PYTHON_ENVIRONMENT.json")
    cuda_bench = load_json("3_FEATURES_CUDA_BENCHMARK_LATEST.json")

    s = ai_index.get("summary", {})
    dq_s = dq.get("summary", {})
    ds_s = ds.get("summary", {})
    wf_s = wf.get("summary", {})
    py_s = pyenv.get("summary", {})
    cb_s = cuda_bench.get("summary", {})
    run_s = run_all.get("summary", {})

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run("ARCHANGEL")
    style_run(run, bold=True, size=24, color="0B2545")

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(10)
    r = subtitle.add_run("Prompt geral v3 | Guia interno de continuidade, calibração e próximos passos")
    style_run(r, size=12, color="555555")

    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(12)
    style_run(meta.add_run(f"Gerado em {datetime.now().strftime('%d/%m/%Y %H:%M')} | Projeto local: {ROOT_DIR}"), size=9, color="555555")

    add_para(
        doc,
        "Este documento é o mapa de continuidade do Archangel. Ele existe para orientar as próximas decisões de código, pesquisa, calibração e execução em ambiente de teste. Deve ser lido como um documento interno de trabalho, não como promessa de performance.",
    )
    add_para(
        doc,
        "Objetivo central: construir um sistema quantitativo realista para operar futuros perpétuos em cripto, inicialmente em testnet/sandbox da Binance, Bybit e Kraken Pro, buscando retorno líquido anual acima de 20% com controle explícito de drawdown, custos, slippage, funding, liquidez e risco operacional.",
    )

    add_h(doc, "1. Onde Estamos Agora", 1)
    add_two_col_table(doc, [
        ("Pipeline geral", f"{run_s.get('status', s.get('pipeline_status'))}; {run_s.get('stages_ok', s.get('stages_ok'))} etapas OK e {run_s.get('stages_error', s.get('stages_error'))} erros no último run completo."),
        ("Universo mapeado", f"{s.get('total_mapped_series')} séries mapeadas; {s.get('features_ok')} séries com features; {s.get('labels_ok')} séries com labels."),
        ("Qualidade de dados", f"{dq_s.get('ml_ready_series_count')} séries ML_READY, {dq_s.get('ml_caution_series_count')} ML_CAUTION e {dq_s.get('ml_blocked_series_count')} ML_BLOCKED."),
        ("Datasets ML", f"{ds_s.get('datasets_ok')} datasets OK; {ds_s.get('total_trainable_rows')} linhas treináveis; todos os datasets atuais estão como ML_CAUTION_ACCEPTABLE."),
        ("Walk-forward", f"{wf_s.get('experiments_ok')} experimentos OK; balanced accuracy média atual de {wf_s.get('avg_balanced_accuracy'):.4f}; backend usado: {wf_s.get('backend_usage')}."),
        ("Ambiente Python/CUDA", f"Stack core, ML e paralela pronta; PyTorch CUDA={py_s.get('cuda_ready_for_pytorch')}; CuPy CUDA={py_s.get('cuda_ready_for_cupy')}; TensorFlow CUDA={py_s.get('cuda_ready_for_tensorflow')}."),
        ("Features com CUDA", f"Benchmark OK; speedup amostral {cb_s.get('speedup_cuda_vs_cpu_total')}x; {cb_s.get('cuda_rolling_accelerated_operations')} operações rolling aceleradas; comparação numérica OK={cb_s.get('numerical_comparison_ok')}."),
    ])

    add_para(
        doc,
        "Leitura honesta: a arquitetura deixou de ser apenas uma ideia e já virou pipeline executável com dados, qualidade, features, labels, datasets e walk-forward. O ponto fraco ainda não é infraestrutura; é transformar predição em PnL líquido realista, com backtest, risco e execução de teste bem separados.",
    )

    add_h(doc, "2. Princípios Que Não Podem Ser Quebrados", 1)
    for item in [
        "Pesquisa não pode alterar execução. O módulo de pesquisa testa ideias; o módulo de execução só aceita estratégias versionadas e aprovadas.",
        "Todo resultado relevante deve terminar em JSON dentro de 0_REGRAS_MANDATO/BASE_JSON, com summary, paths, schema_version e run_id quando aplicável.",
        "Toda feature usada por modelo deve ser auditável e livre de vazamento. Features permitidas: feat_*, xasset_* e regime_*; labels e metadados nunca entram como input direto.",
        "A meta de 20% líquido ao ano só é válida depois de custos, slippage, funding, liquidez, latência, partial fills e drawdown.",
        "Acurácia isolada não decide estratégia. O que importa é retorno líquido ajustado ao risco, estabilidade fora da amostra e sobrevivência sob stress test.",
        "Alavancagem deve ser consequência do orçamento de risco. Limite máximo desejado: 5x, com haircut e kill switch.",
    ]:
        add_bullet(doc, item)

    add_h(doc, "3. Arquitetura-Alvo", 1)
    add_status_table(doc, [
        ("Research layer", "Parcialmente pronta", "Overfitting e excesso de combinações", "Criar registry de experimentos e filas de hipóteses."),
        ("Data layer", "Pronta e auditada", "TIME_GAPS em parte das séries", "Priorizar limpeza/triagem das séries que entram no MVP."),
        ("Feature/ML layer", "Operacional", "Features convencionais ainda dominam", "Expandir features não convencionais com benchmark e ablação."),
        ("Labels", "Operacional", "Label pode não capturar execução real", "Testar variações triple barrier, meta-labeling e labels líquidos."),
        ("Walk-forward", "Operacional com CUDA", "Balanced accuracy ainda modesta", "Adicionar purging/embargo mais explícito e salvar análises por regime."),
        ("Backtest layer", "Próximo gargalo", "PnL sem execução realista engana", "Construir 7_BACKTEST_PORTFOLIO.py robusto."),
        ("Risk engine", "A desenhar", "Drawdown e alavancagem", "Criar limites de exposição, sizing, stop global e kill switch."),
        ("Execution layer", "Ainda não ligar live", "Risco operacional", "Criar execução apenas em testnet/sandbox."),
        ("Monitoring", "Pendente", "Sem reconciliação contínua", "Criar logs, alertas e reconciliação exchange vs sistema."),
    ])

    add_h(doc, "4. Próximos Passos Prioritários", 1)
    add_number(doc, "Ligar o módulo de teste de execução em ambiente de teste para Bybit, Binance e Kraken Pro. Nada de live nesta fase.")
    add_number(doc, "Construir ou completar o backtest de portfólio com custos, funding, slippage, stops, take profit, sizing, drawdown e métricas líquidas.")
    add_number(doc, "Criar um registry de experimentos para versionar dataset, features, labels, modelo, hiperparâmetros, custos, risco, métricas e status.")
    add_number(doc, "Expandir features não convencionais, mas sempre com ablação, walk-forward e stress de custos.")
    add_number(doc, "Depois de cripto perp estar estável, adaptar o pipeline para futuros de commodities, começando por coffee, cocoa e cotton, com primeiro e segundo vencimento.")

    add_h(doc, "5. Execução em Teste: Bybit, Binance e Kraken Pro", 1)
    add_para(doc, "A próxima fase deve ser paper/test execution. O objetivo não é ganhar dinheiro ainda; é provar que o sistema consegue gerar sinal, transformar sinal em ordem simulada, reconciliar posição e registrar tudo sem ambiguidade.")
    add_two_col_table(doc, [
        ("Escopo inicial", "Perpétuos lineares líquidos: BTCUSDT, ETHUSDT, SOLUSDT, BNBUSDT e XRPUSDT, salvo restrição de testnet."),
        ("Exchanges", "Bybit, Binance e Kraken Pro, sempre em ambiente de teste/sandbox/testnet quando disponível."),
        ("Ordens", "Começar com market/limit simples, reduce-only, cancelamento controlado e sem pirâmide de posição."),
        ("Logs obrigatórios", "signal_id, model_version, dataset_version, risk_config, order_id, exchange_order_id, fill, slippage estimado e realizado."),
        ("Kill switch", "Parar novas ordens em caso de erro de API, perda de conexão, drawdown diário, divergência de posição ou latência acima do limite."),
        ("Critério de sucesso", "Rodar por pelo menos 30 a 60 dias em teste com estabilidade operacional e diferença aceitável entre sinal esperado e execução simulada."),
    ])

    add_h(doc, "6. Features Não Convencionais", 1)
    add_para(doc, "O sistema deve continuar usando features clássicas como baseline, mas o diferencial para um trader individual tende a vir de combinações que grandes competidores ignoram ou descartam por escala, fricção ou granularidade.")
    for item in [
        "Timeframes fora do consenso: 7m, 13m, 23m, 37m e combinações entre timeframe de decisão e timeframe de contexto.",
        "Regime features: compressão/expansão de volatilidade, persistência de range, velocidade de mudança de regime e assimetria entre alta e queda.",
        "Cross-asset features: liderança/atraso entre BTC, ETH, SOL, BNB e XRP; beta/correlação rolling; resíduo contra fator de mercado cripto.",
        "Microestrutura quando disponível: intensidade de trades, agressor imbalance, taker buy/sell volume, proxy de spread e variação de liquidez intradiária.",
        "Features de custo: fee drag esperado, funding drag, spread estimado, sensibilidade a slippage e distância entre edge previsto e custo total.",
        "Features de robustez: sinal só é elegível quando múltiplas famílias independentes concordam, evitando depender de um único indicador.",
        "Meta-features: estabilidade recente do modelo, calibração de probabilidade, erro recente por regime e degradação fora da amostra.",
    ]:
        add_bullet(doc, item)

    add_h(doc, "7. Labels e Target", 1)
    add_para(doc, "O target principal continua sendo Triple Barrier, porque conversa com trading real: take profit, stop loss e tempo máximo. Porém o label precisa evoluir para refletir custo líquido e execução.")
    add_two_col_table(doc, [
        ("Label base", "Triple barrier: +1 para take profit primeiro, -1 para stop primeiro, 0 para nenhum evento até o horizonte."),
        ("Versão líquida", "Barreiras calculadas depois de fee, slippage e funding estimado. Sem isso, o modelo aprende edge bruto que pode desaparecer na execução."),
        ("Meta-labeling", "Depois de um sinal primário, treinar um segundo modelo para estimar se vale executar aquele sinal específico."),
        ("Long/short separado", "Separar labels e métricas de long e short para evitar misturar regimes assimétricos."),
        ("Overlapping labels", "Usar purging/embargo no walk-forward para reduzir vazamento estatístico entre janelas sobrepostas."),
    ])

    add_h(doc, "8. Backtest e Risk Engine", 1)
    add_para(doc, "A etapa que mais vai aproximar o projeto da realidade é o backtest com motor de risco. O modelo não deve ser avaliado por acurácia; ele deve ser avaliado por PnL líquido, drawdown e estabilidade.")
    for item in [
        "Entrada: probabilidade calibrada, margem sobre custo total, filtro de regime e filtro de liquidez.",
        "Saída: stop fixo, stop por volatilidade, stop por tempo, trailing stop e queda da probabilidade do modelo.",
        "Sizing: volatilidade alvo, risco máximo por trade, exposição máxima por ativo, exposição máxima por exchange e alavancagem máxima de 5x.",
        "Drawdown control: reduzir risco após drawdown diário, semanal e mensal; pausar novas entradas em stress.",
        "Métricas: CAGR líquido, max drawdown, Calmar, Sharpe, Sortino, profit factor, turnover, fee drag, funding drag, slippage drag e % de janelas OOS positivas.",
    ]:
        add_bullet(doc, item)

    add_h(doc, "9. Commodities: Coffee, Cocoa e Cotton", 1)
    add_para(doc, "A expansão para commodities deve vir depois de cripto perp estar estável, porque futuros de commodities exigem cuidado com vencimento, rolagem, curva, calendário e liquidez.")
    add_two_col_table(doc, [
        ("Séries iniciais", "Primeiro e segundo futuro de coffee, cocoa e cotton, com preços diários de ajuste."),
        ("Features centrais", "Retorno do contrato, spread primeiro-segundo, inclinação da curva, basis quando disponível, carry, sazonalidade e volatilidade realizada."),
        ("Regressores", "Dólar, rates, energia, índices agrícolas, commodities correlatas, clima quando houver fonte confiável e calendário de safra."),
        ("Ajuste de dados", "Tratar rolagem, vencimento, liquidez do contrato, feriados e mudanças de especificação."),
        ("Labels", "Targets por horizonte diário/semanal, triple barrier adaptado à volatilidade diária e custos de rolagem."),
    ])

    add_h(doc, "10. Critérios Realistas de Sucesso", 1)
    add_two_col_table(doc, [
        ("Meta de pesquisa", "Encontrar estratégias com retorno líquido anualizado acima de 20%, max drawdown controlado e estabilidade fora da amostra."),
        ("Mínimo aceitável", "Calmar acima de 1, Sharpe acima de 1.2, OOS positivo em mais de 60% das janelas e queda controlada sob custos maiores."),
        ("Teste operacional", "30 a 60 dias em testnet sem divergência material entre sinal, ordem, fill e posição reconciliada."),
        ("Live futuro", "Somente com capital pequeno, kill switch ativo, limites de perda e monitoramento contínuo."),
        ("Sinal vermelho", "Se o retorno só aparece com custo baixo, alavancagem alta, janela curta ou ativo específico, tratar como provável overfit."),
    ])

    add_h(doc, "11. Checklist de Continuidade", 1)
    for item in [
        "Rerodar etapa 3 com CUDA forçada apenas em teste controlado: ARCHANGEL_FEATURE_CUDA_MODE=cuda e ARCHANGEL_FEATURE_CUDA_MAX_WORKERS=2.",
        "Comparar o run completo da etapa 3 com o benchmark atual antes de tornar CUDA padrão.",
        "Conferir se 3_FEATURES_CUDA_BENCHMARK_LATEST.json continua OK após qualquer mudança em features.",
        "Garantir que toda nova etapa grave JSON útil em BASE_JSON e seja incluída no ARCHANGEL_AI_CONTEXT_INDEX.json.",
        "Criar módulo de execução testnet com dry-run, logs e kill switch antes de qualquer ordem real.",
        "Implementar backtest de portfólio antes de aumentar o universo de modelos.",
        "Adicionar novas features em blocos pequenos: hipótese, implementação, benchmark, ablação, walk-forward, stress de custos.",
        "Só considerar alavancagem depois que a estratégia sem alavancagem sobreviver ao custo líquido e ao walk-forward.",
    ]:
        add_bullet(doc, item)

    add_h(doc, "12. Prompt Operacional Para Próximas Sessões Com Codex", 1)
    add_para(doc, "Sempre que o projeto for recalibrado, começar lendo este documento e depois o índice mestre:")
    add_bullet(doc, str(BASE_JSON_DIR / "ARCHANGEL_AI_CONTEXT_INDEX.json"))
    add_para(doc, "Pedido padrão para o Codex:")
    add_para(
        doc,
        "Leia o ARCHANGEL_AI_CONTEXT_INDEX.json, confira os JSONs *_LATEST no BASE_JSON, identifique o estado atual do pipeline, preserve governança anti-leakage e proponha a próxima mudança incremental com teste, benchmark e relatório JSON. Não pule para live trading; execução deve ser apenas em testnet/sandbox até aprovação explícita.",
    )

    add_h(doc, "13. Decisão Atual", 1)
    add_para(
        doc,
        "O Archangel está pronto para sair da fase de pipeline ML básico e entrar na fase de execução simulada/testnet e backtest realista. A prioridade agora é transformar previsões em decisões operáveis com risco controlado, não aumentar cegamente número de modelos.",
    )
    add_para(
        doc,
        "A meta de retorno líquido acima de 20% ao ano é possível como objetivo de pesquisa, mas só deve ser considerada séria se sobreviver a custos conservadores, períodos ruins, múltiplos ativos, walk-forward e execução de teste. O projeto deve preferir robustez a brilho de backtest.",
    )

    section = doc.add_section(WD_SECTION.CONTINUOUS)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_run(footer.add_run("ARCHANGEL | Guia interno v3 | uso pessoal"), size=8, color="555555")

    doc.save(OUTPUT_PATH)
    return OUTPUT_PATH


if __name__ == "__main__":
    path = build_doc()
    print(path)
