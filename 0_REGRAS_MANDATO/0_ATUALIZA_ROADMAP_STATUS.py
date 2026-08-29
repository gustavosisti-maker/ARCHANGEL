# -*- coding: utf-8 -*-
from __future__ import annotations

import json
import os
import subprocess
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from openpyxl import Workbook
from openpyxl.chart import BarChart, Reference
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter

SCRIPT_NAME = '0_ATUALIZA_ROADMAP_STATUS.py'
SCHEMA_VERSION = 'ARCHANGEL_ROADMAP_STATUS_REPORT_1.0'
ROOT_DIR = Path(__file__).resolve().parent.parent
RULES_DIR = ROOT_DIR / '0_REGRAS_MANDATO'
BASE_JSON_DIR = RULES_DIR / 'BASE_JSON'
DOCX_PATH = RULES_DIR / 'PROMPT_GERAL_v3_ARCHANGEL_ROADMAP_INTERNO.docx'
XLSX_PATH = RULES_DIR / 'ARCHANGEL_STATUS_PARAMETROS_RESULTADOS.xlsx'
REPORT_JSON_PATH = BASE_JSON_DIR / 'ROADMAP_STATUS_REPORT_LATEST.json'

JSONS = [
    '00_RUN_ALL_RUN_REPORT_LATEST.json','ARCHANGEL_AI_CONTEXT_INDEX.json','RUN_STATE.json',
    'ARCHANGEL_MACHINE_PROFILE.json','ARCHANGEL_PYTHON_ENVIRONMENT.json','MAPA_ATIVOS.json',
    'DATA_QUALITY_REPORT.json','DATA_QUALITY_ROOT_CAUSE_REPORT.json','3_JSON_FEATURES.json',
    '3_FEATURES_RUN_REPORT_LATEST.json','3_FEATURES_CUDA_BENCHMARK_LATEST.json','3_FEATURES_RETRY_PLAN_LATEST.json',
    '4_JSON_LABELS.json','4_LABELS_RUN_REPORT_LATEST.json','5_JSON_DATASETS_ML.json','5_DATASETS_ML_RUN_REPORT_LATEST.json',
    '6_JSON_WALK_FORWARD.json','6_WALK_FORWARD_RUN_REPORT_LATEST.json','7_JSON_BACKTEST_PORTFOLIO.json',
    '7_BACKTEST_PORTFOLIO_RUN_REPORT_LATEST.json','7_BACKTEST_PARAM_SEARCH_LATEST.json','7_BACKTEST_VALIDATION_LATEST.json',
    '7_BACKTEST_STRESS_LATEST.json','7_EXPERIMENT_REGISTRY_LATEST.json','8_JSON_EXPERIMENT_REGISTRY.json',
    '8_EXPERIMENT_REGISTRY_LATEST.json','8_EXPERIMENT_REGISTRY_RUN_REPORT_LATEST.json','8_EXPERIMENT_REGISTRY_VALIDATION_LATEST.json',
    'COST_MODEL.json'
]


def load_json(name: str) -> dict[str, Any]:
    path = BASE_JSON_DIR / name
    if not path.is_file():
        return {}
    last = None
    for _ in range(5):
        try:
            data = json.loads(path.read_text(encoding='utf-8'))
            return data if isinstance(data, dict) else {}
        except Exception as exc:
            last = exc
    return {'_read_error': str(last)}


def nested(data: dict[str, Any], *keys: str, default: Any = None) -> Any:
    cur: Any = data
    for key in keys:
        if not isinstance(cur, dict):
            return default
        cur = cur.get(key)
    return default if cur is None else cur


def fnum(value: Any, default: float | None = None) -> float | None:
    try:
        out = float(value)
    except Exception:
        return default
    return out if out == out else default


def fint(value: Any) -> str:
    try:
        return f'{int(value):,}'.replace(',', '.')
    except Exception:
        return 'n/d'


def fpct(value: Any) -> str:
    val = fnum(value)
    return 'n/d' if val is None else f'{val * 100:.2f}%'


def fdec(value: Any) -> str:
    val = fnum(value)
    return 'n/d' if val is None else f'{val:.4f}'


def inventory() -> list[dict[str, Any]]:
    rows = []
    for name in JSONS:
        path = BASE_JSON_DIR / name
        data = load_json(name)
        system = data.get('system', {}) if isinstance(data.get('system'), dict) else {}
        rows.append({
            'file': name,
            'exists': path.is_file(),
            'size_bytes': path.stat().st_size if path.is_file() else None,
            'mtime_local': datetime.fromtimestamp(path.stat().st_mtime).isoformat(timespec='seconds') if path.is_file() else None,
            'schema_version': data.get('schema_version'),
            'run_id': data.get('run_id') or system.get('run_id'),
            'generated_at': data.get('generated_at') or data.get('generated_at_utc') or system.get('generated_at') or system.get('generated_at_utc'),
            'read_error': data.get('_read_error'),
        })
    return rows


def mtime(name: str) -> float:
    path = BASE_JSON_DIR / name
    return path.stat().st_mtime if path.is_file() else 0.0


def assess() -> dict[str, Any]:
    data = {name: load_json(name) for name in JSONS}
    inv = inventory()
    run = data['00_RUN_ALL_RUN_REPORT_LATEST.json']
    mapa = data['MAPA_ATIVOS.json']; dq = data['DATA_QUALITY_REPORT.json']
    feat = data['3_JSON_FEATURES.json']; feat_run = data['3_FEATURES_RUN_REPORT_LATEST.json']
    labels = data['4_JSON_LABELS.json']; ds = data['5_JSON_DATASETS_ML.json']; wf = data['6_JSON_WALK_FORWARD.json']
    bt = data['7_JSON_BACKTEST_PORTFOLIO.json']; val = data['7_BACKTEST_VALIDATION_LATEST.json']; stress = data['7_BACKTEST_STRESS_LATEST.json']
    reg = data['8_EXPERIMENT_REGISTRY_LATEST.json']; pyenv = data['ARCHANGEL_PYTHON_ENVIRONMENT.json']; bench = data['3_FEATURES_CUDA_BENCHMARK_LATEST.json']
    run_s = run.get('summary', {}); dq_s = dq.get('summary', {}); feat_s = feat.get('summary', {}); feat_run_s = feat_run.get('summary', {})
    label_s = labels.get('summary', {}); ds_s = ds.get('summary', {}); wf_s = wf.get('summary', {}); bt_s = bt.get('summary', {}); reg_s = reg.get('summary', {})
    run_time = mtime('00_RUN_ALL_RUN_REPORT_LATEST.json'); feat_time = max(mtime('3_JSON_FEATURES.json'), mtime('3_FEATURES_RUN_REPORT_LATEST.json'))
    downstream = {'labels': mtime('4_JSON_LABELS.json'), 'datasets': mtime('5_JSON_DATASETS_ML.json'), 'walk_forward': mtime('6_JSON_WALK_FORWARD.json'), 'backtest': mtime('7_JSON_BACKTEST_PORTFOLIO.json'), 'registry': mtime('8_EXPERIMENT_REGISTRY_LATEST.json')}
    stale = [k for k, v in downstream.items() if feat_time and v and v + 30 < feat_time]
    newer = [r['file'] for r in inv if (BASE_JSON_DIR / r['file']).is_file() and (BASE_JSON_DIR / r['file']).stat().st_mtime > run_time + 30]
    cagr = fnum(bt_s.get('portfolio_cagr')); total_ret = fnum(bt_s.get('portfolio_total_return')); max_dd = fnum(bt_s.get('portfolio_max_drawdown'))
    bal = fnum(wf_s.get('avg_balanced_accuracy'))
    red = []
    if newer: red.append('Há JSONs mais novos que o último relatório do 00_RUN_ALL; este snapshot pode ser parcial ou refletir run ainda em andamento.')
    if stale: red.append('A etapa 3/features está mais nova que downstreams; labels/datasets/walk-forward/backtest/registry podem estar defasados até o run atual terminar.')
    if cagr is not None and cagr < 0: red.append('Backtest de portfólio atual tem CAGR negativo; não existe evidência econômica para testnet por performance.')
    if max_dd is not None and abs(max_dd) > 0.08: red.append('Max drawdown observado ultrapassa a referência de 8%; isto bloqueia qualquer leitura otimista.')
    if bal is not None and bal < 0.55: red.append('Balanced accuracy média está apenas levemente acima de 50%; edge estatístico ainda é fraco.')
    if dq_s.get('ml_caution_series_count', 0): red.append('Muitas séries seguem em ML_CAUTION, principalmente por TIME_GAPS; isso contamina confiança nos resultados.')
    verdict = 'SNAPSHOT PARCIAL: AGUARDAR FIM DO 00_RUN_ALL ANTES DE DECIDIR' if newer else 'PESQUISA AVANÇOU, MAS NÃO HÁ SINAL ECONÔMICO APROVADO'
    stages = [
        ['Git/GitHub','OK','Repo local ligado ao origin/main e CI leve passou.','Repo é público; dados/JSONs/docx continuam fora.','Commits pequenos antes de mudanças grandes.'],
        ['Dados','PARCIAL','{} séries mapeadas; {} OHLCV.'.format(fint(nested(mapa,'summary','total_series_parquet_mapped')), fint(nested(mapa,'summary','total_ohlcv_series'))),'Warnings e séries unknown ainda existem.','Isolar universo MVP.'],
        ['Qualidade','USÁVEL COM CAUTELA','{} ML_READY, {} ML_CAUTION, {} ML_BLOCKED.'.format(fint(dq_s.get('ml_ready_series_count')), fint(dq_s.get('ml_caution_series_count')), fint(dq_s.get('ml_blocked_series_count'))),'TIME_GAPS é raiz dominante.','Priorizar séries realmente usadas.'],
        ['Features','OPERACIONAL','{} séries OK; {} features; backend atual {}.'.format(fint(feat_s.get('series_ok')), fint(feat_s.get('feature_catalog_count')), nested(feat_run_s,'feature_compute_backend','series_backend_counts')),'Último run de features aparece em CPU/pandas.','CUDA só com benchmark e fallback.'],
        ['Labels','OPERACIONAL/DEFASAGEM POSSÍVEL','{} séries; {} labels.'.format(fint(label_s.get('series_ok')), fint(label_s.get('labels_catalog_count'))),'Pode estar defasado se features acabou de atualizar.','Confirmar timestamps após o run.'],
        ['Datasets ML','OPERACIONAL, MAS CAUTION','{} datasets; {} linhas treináveis.'.format(fint(ds_s.get('datasets_ok')), fint(ds_s.get('total_trainable_rows'))),'Todos atuais são ML_CAUTION_ACCEPTABLE.','Medir robustez por causa raiz.'],
        ['Walk-forward','TÉCNICO OK','{} experimentos; balanced accuracy {}.'.format(fint(wf_s.get('experiments_ok')), fdec(wf_s.get('avg_balanced_accuracy'))),'Métrica fraca para edge operacional.','Usar como filtro, não aprovação.'],
        ['Backtest portfolio','MOTOR OK, RESULTADO RUIM','{} trades; CAGR {}; retorno {}; DD {}; validação {}.'.format(fint(bt_s.get('portfolio_total_trades')), fpct(cagr), fpct(total_ret), fpct(max_dd), nested(val,'summary','status')),'PnL líquido negativo e DD acima de 8%.','Usar para ablação/stress, não testnet.'],
        ['Stress','OK COMO TESTE','Pior cenário {}; retorno {}; DD {}.'.format(nested(stress,'summary','worst_scenario'), fpct(nested(stress,'summary','worst_total_return')), fpct(nested(stress,'summary','worst_max_drawdown'))),'Stress confirma fragilidade econômica atual.','Comparar novas features sob stress.'],
        ['Experiment registry','FORMALIZADO','{} experimentos; validação {}; status {}.'.format(fint(reg_s.get('registry_rows')), reg_s.get('validation_status'), reg_s.get('status')),'Registry registra evidência, não aprova execução.','Virar tabela mestre de pesquisa.'],
        ['CUDA','PRONTO CONTROLADO','PyTorch CUDA={}; CuPy CUDA={}; benchmark numérico OK={}.'.format(nested(pyenv,'summary','cuda_ready_for_pytorch'), nested(pyenv,'summary','cuda_ready_for_cupy'), nested(bench,'summary','numerical_comparison_ok')),'Acelerar sem equivalência numérica cria falso progresso.','Migrar apenas blocos caros.'],
    ]
    roadmap = [
        [1,'Esperar o 00_RUN_ALL atual terminar e reavaliar timestamps','Snapshot atual pode estar parcial.','AGORA'],
        [2,'Atualizar DOCX/XLSX/JSON automaticamente ao fim do 00_RUN_ALL','Evita pedir prompt reverso manual.','IMPLEMENTADO'],
        [3,'Criar fila de hipóteses e ablação formal de features','O resultado econômico atual é ruim.','PRÓXIMO'],
        [4,'Expandir features não convencionais em blocos pequenos','Edge precisa ser medido contra baseline e custos.','PRÓXIMO'],
        [5,'Criar módulo testnet Bybit/Binance/Kraken Pro','Só depois de estratégia candidata sobreviver ao backtest/stress.','DEPOIS'],
        [6,'Adaptar para futuros de commodities: coffee, cocoa, cotton','Exige vencimento, rolagem, curva, calendário e liquidez.','FUTURO'],
    ]
    params = [
        ['target_annual_return_min', bt_s.get('target_annual_return_min'), '7_JSON_BACKTEST_PORTFOLIO', 'Alvo econômico mínimo; não é aprovação automática.'],
        ['reference_drawdown_limit', bt_s.get('reference_drawdown_limit'), '7_JSON_BACKTEST_PORTFOLIO', 'Referência temporária de 8%.'],
        ['approval_status', bt_s.get('approval_status'), '7_JSON_BACKTEST_PORTFOLIO', 'Pesquisa separada de aprovação.'],
        ['feature_backend_mode', nested(feat_run_s,'feature_compute_backend','mode'), '3_FEATURES_RUN_REPORT', 'Último run de features em auto.'],
        ['walk_forward_backend_usage', wf_s.get('backend_usage'), '6_JSON_WALK_FORWARD', 'Treino já usa torch_cuda.'],
        ['cuda_ready_for_pytorch', nested(pyenv,'summary','cuda_ready_for_pytorch'), 'ARCHANGEL_PYTHON_ENVIRONMENT', 'Pronto para uso controlado.'],
        ['cuda_ready_for_cupy', nested(pyenv,'summary','cuda_ready_for_cupy'), 'ARCHANGEL_PYTHON_ENVIRONMENT', 'Pronto para blocos selecionados.'],
    ]
    return {'schema_version': SCHEMA_VERSION, 'system': {'name': 'ARCHANGEL', 'layer': 'ROADMAP_STATUS', 'script': SCRIPT_NAME, 'generated_at_utc': datetime.now(timezone.utc).isoformat(timespec='seconds')}, 'paths': {'docx_path': str(DOCX_PATH), 'xlsx_path': str(XLSX_PATH), 'report_json_path': str(REPORT_JSON_PATH), 'base_json_dir': str(BASE_JSON_DIR)}, 'summary': {'verdict': verdict, 'likely_run_in_progress_or_partial_snapshot': bool(newer), 'downstream_stale_after_features': stale, 'red_flags_count': len(red), 'pipeline_status_latest_completed': run_s.get('status'), 'features_ok': feat_s.get('series_ok'), 'labels_ok': label_s.get('series_ok'), 'datasets_ok': ds_s.get('datasets_ok'), 'walk_forward_experiments_ok': wf_s.get('experiments_ok'), 'walk_forward_avg_balanced_accuracy': wf_s.get('avg_balanced_accuracy'), 'portfolio_cagr': cagr, 'portfolio_total_return': total_ret, 'portfolio_max_drawdown': max_dd, 'formal_registry_status': reg_s.get('status'), 'formal_registry_rows': reg_s.get('registry_rows')}, 'staleness': {'jsons_newer_than_last_run_all_report': newer, 'downstream_stale_after_features': stale}, 'red_flags': red, 'stage_assessment': [{'stage':r[0], 'status':r[1], 'evidence':r[2], 'risk':r[3], 'next_action':r[4]} for r in stages], 'roadmap': [{'priority':r[0], 'step':r[1], 'why':r[2], 'status':r[3]} for r in roadmap], 'parameters': [{'parameter':r[0], 'value':r[1], 'source':r[2], 'note':r[3]} for r in params], 'json_inventory': inv}


def write_json(path: Path, payload: dict[str, Any]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    tmp = path.with_suffix(path.suffix + '.tmp')
    tmp.write_text(json.dumps(payload, ensure_ascii=False, indent=2, default=str), encoding='utf-8')
    os.replace(tmp, path)


def set_font(run, bold=False, size=10.5, color='222222'):
    run.bold = bold; run.font.size = Pt(size); run.font.color.rgb = RGBColor.from_string(color); run.font.name = 'Calibri'
    run._element.get_or_add_rPr().rFonts.set(qn('w:ascii'), 'Calibri'); run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')


def shade(cell, fill):
    pr = cell._tc.get_or_add_tcPr(); shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), fill); pr.append(shd)


def table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER; table.autofit = False
    for row in table.rows:
        for i, w in enumerate(widths):
            if i < len(row.cells): row.cells[i].width = Inches(w); row.cells[i].vertical_alignment = WD_ALIGN_VERTICAL.TOP


def para(doc, text, bold_prefix=None, color='222222'):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(6); p.paragraph_format.line_spacing = 1.15
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix); set_font(r, True, 10.5, color); r = p.add_run(text[len(bold_prefix):]); set_font(r, False, 10.5, color)
    else: set_font(p.add_run(text), False, 10.5, color)


def bullet(doc, text, color='222222'):
    p = doc.add_paragraph(style='List Bullet'); p.paragraph_format.space_after = Pt(4); set_font(p.add_run(text), False, 10, color)


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level); p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(5)
    for r in p.runs: set_font(r, True, 16 if level == 1 else 12.5, '0B2545' if level == 1 else '1F4D78')


def matrix(doc, headers, rows, widths):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = 'Table Grid'
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h; shade(t.rows[0].cells[i], 'E8EEF5')
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row): cells[i].text = str(value)
    for row in t.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for r in p.runs: set_font(r, p is t.rows[0].cells[0].paragraphs[0], 8.8)
    table_width(t, widths)


def build_docx(report):
    doc = Document(); sec = doc.sections[0]; sec.top_margin = Inches(0.8); sec.bottom_margin = Inches(0.75); sec.left_margin = Inches(0.85); sec.right_margin = Inches(0.85)
    normal = doc.styles['Normal']; normal.font.name = 'Calibri'; normal.font.size = Pt(10.5)
    title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.LEFT; set_font(title.add_run('ARCHANGEL'), True, 24, '0B2545')
    sub = doc.add_paragraph(); set_font(sub.add_run('Roadmap interno v3 | Prompt reverso, status real e próximos passos'), False, 12, '555555')
    meta = doc.add_paragraph(); set_font(meta.add_run(f"Atualizado em {datetime.now().strftime('%d/%m/%Y %H:%M')} | Projeto local: {ROOT_DIR}"), False, 9, '555555')
    heading(doc, '1. Veredito Executivo')
    matrix(doc, ['Campo','Valor'], [['Veredito', report['summary']['verdict']], ['Snapshot', 'Parcial/em andamento' if report['summary']['likely_run_in_progress_or_partial_snapshot'] else 'Coerente com último run completo'], ['Meta econômica','Retorno líquido mínimo desejado acima de 20% ao ano; ainda não demonstrado.'], ['Regra','Pesquisa não é aprovação de execução.']], [1.8,4.7])
    heading(doc, '2. Alertas Que Não Devem Ser Ignorados')
    for r in report['red_flags']: bullet(doc, r, '9B1C1C')
    heading(doc, '3. Estado Atual Por Camada')
    matrix(doc, ['Camada','Estado','Evidência','Risco','Próxima ação'], [[x['stage'],x['status'],x['evidence'],x['risk'],x['next_action']] for x in report['stage_assessment']], [1.05,1.05,1.8,1.3,1.3])
    heading(doc, '4. Métricas Duras')
    s=report['summary']; matrix(doc, ['Métrica','Valor'], [['Features OK',fint(s.get('features_ok'))],['Labels OK',fint(s.get('labels_ok'))],['Datasets OK',fint(s.get('datasets_ok'))],['Walk-forward OK',fint(s.get('walk_forward_experiments_ok'))],['Balanced accuracy média',fdec(s.get('walk_forward_avg_balanced_accuracy'))],['Backtest CAGR',fpct(s.get('portfolio_cagr'))],['Backtest retorno total',fpct(s.get('portfolio_total_return'))],['Backtest max drawdown',fpct(s.get('portfolio_max_drawdown'))],['Registry formal',f"{s.get('formal_registry_status')} com {fint(s.get('formal_registry_rows'))} experimentos"]], [2.1,4.4])
    para(doc, 'Leitura cruelmente realista: a engenharia do pipeline está avançada, mas o resultado econômico atual é ruim. A próxima vitória é descobrir, via ablação e stress, se existe edge líquido robusto.', 'Leitura cruelmente realista:')
    heading(doc, '5. Roadmap Recomendado')
    for r in report['roadmap']: bullet(doc, f"{r['priority']}. {r['step']} [{r['status']}]. Motivo: {r['why']}")
    heading(doc, '6. Como Usar Este Documento Com O Codex')
    para(doc, 'Leia este DOCX, depois ROADMAP_STATUS_REPORT_LATEST.json, ARCHANGEL_AI_CONTEXT_INDEX.json e o Excel de status. Peça sempre diagnóstico objetivo, preservando anti-leakage e separando pesquisa de aprovação.')
    heading(doc, '7. Coerência Dos JSONs')
    matrix(doc, ['Campo','Valor'], [['Downstreams defasados', ', '.join(report['staleness']['downstream_stale_after_features']) or 'nenhum'], ['JSONs mais novos que run_all', len(report['staleness']['jsons_newer_than_last_run_all_report'])]], [2.1,4.4])
    doc.save(DOCX_PATH)


def style_sheet(ws):
    ws.sheet_view.showGridLines = False; thin = Side(style='thin', color='DADCE0')
    for row in ws.iter_rows():
        for c in row:
            c.font = Font(name='Calibri', size=10, color='222222'); c.alignment = Alignment(vertical='top', wrap_text=True); c.border = Border(bottom=thin)
            if c.row == 1: c.font = Font(name='Calibri', size=10, bold=True, color='FFFFFF'); c.fill = PatternFill('solid', fgColor='0B2545')
    ws.freeze_panes = 'A2'; ws.auto_filter.ref = ws.dimensions
    for i in range(1, ws.max_column+1):
        width = 10
        for j in range(1, min(ws.max_row,250)+1):
            v = ws.cell(j,i).value
            if v is not None: width = max(width, min(60, len(str(v))+2))
        ws.column_dimensions[get_column_letter(i)].width = width


def rowsheet(wb, name, headers, rows):
    ws = wb.create_sheet(name); ws.append(headers)
    for r in rows: ws.append([json.dumps(r.get(h), ensure_ascii=False, default=str) if isinstance(r.get(h),(dict,list)) else r.get(h) for h in headers])
    style_sheet(ws); return ws


def build_xlsx(report):
    wb = Workbook(); ws = wb.active; ws.title = 'Dashboard'; ws.append(['Campo','Valor'])
    for k,v in report['summary'].items(): ws.append([k, json.dumps(v,ensure_ascii=False,default=str) if isinstance(v,(dict,list)) else v])
    ws['E1']='Métrica'; ws['F1']='Valor'
    metrics=[('Balanced Accuracy',report['summary']['walk_forward_avg_balanced_accuracy']),('CAGR',report['summary']['portfolio_cagr']),('Total Return',report['summary']['portfolio_total_return']),('Max Drawdown',report['summary']['portfolio_max_drawdown'])]
    for i,(k,v) in enumerate(metrics,2): ws.cell(i,5).value=k; ws.cell(i,6).value=v
    chart=BarChart(); chart.title='Métricas principais'; chart.add_data(Reference(ws,min_col=6,min_row=1,max_row=5),titles_from_data=True); chart.set_categories(Reference(ws,min_col=5,min_row=2,max_row=5)); chart.height=7; chart.width=14; ws.add_chart(chart,'H2')
    style_sheet(ws)
    rowsheet(wb,'Pipeline_Status',['stage','status','evidence','risk','next_action'],report['stage_assessment'])
    rowsheet(wb,'Roadmap',['priority','step','why','status'],report['roadmap'])
    rowsheet(wb,'Parameters',['parameter','value','source','note'],report['parameters'])
    rowsheet(wb,'JSON_Inventory',['file','exists','size_bytes','mtime_local','schema_version','run_id','generated_at','read_error'],report['json_inventory'])
    flags = wb.create_sheet('Red_Flags'); flags.append(['red_flag']); [flags.append([x]) for x in report['red_flags']]; style_sheet(flags)
    wb.save(XLSX_PATH)


def update_ai_index(report):
    path = BASE_JSON_DIR/'ARCHANGEL_AI_CONTEXT_INDEX.json'
    if not path.is_file(): return
    idx = load_json('ARCHANGEL_AI_CONTEXT_INDEX.json')
    idx.setdefault('summary', {})['roadmap_status_verdict'] = report['summary']['verdict']
    idx['summary']['roadmap_status_generated_at_utc'] = report['system']['generated_at_utc']
    order = idx.setdefault('ai_reading_order', [])
    if 'ROADMAP_STATUS_REPORT_LATEST.json' not in order: order.insert(1, 'ROADMAP_STATUS_REPORT_LATEST.json')
    write_json(path, idx)


def main() -> int:
    BASE_JSON_DIR.mkdir(parents=True, exist_ok=True)
    report = assess(); write_json(REPORT_JSON_PATH, report); build_docx(report); build_xlsx(report); update_ai_index(report)
    print('ARCHANGEL roadmap/status atualizado')
    print(f'DOCX: {DOCX_PATH}')
    print(f'XLSX: {XLSX_PATH}')
    print(f'JSON: {REPORT_JSON_PATH}')
    print(f"Veredito: {report['summary']['verdict']}")
    return 0

if __name__ == '__main__':
    raise SystemExit(main())
