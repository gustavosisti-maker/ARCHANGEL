# -*- coding: utf-8 -*-
from __future__ import annotations

import json
import os
from collections import Counter
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from openpyxl import Workbook
from openpyxl.chart import BarChart, Reference
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.worksheet.table import Table, TableStyleInfo
from openpyxl.utils import get_column_letter

SCRIPT_NAME = '09_ATUALIZA_ROADMAP_STATUS.py'
SCHEMA_VERSION = 'ARCHANGEL_ROADMAP_STATUS_REPORT_2.0'
ROOT_DIR = Path(__file__).resolve().parent.parent
RULES_DIR = ROOT_DIR / '0_REGRAS_MANDATO'
BASE_JSON_DIR = RULES_DIR / 'BASE_JSON'
DOCX_PATH = RULES_DIR / 'PROMPT_GERAL_v3_ARCHANGEL_ROADMAP_INTERNO.docx'
XLSX_PATH = RULES_DIR / 'ARCHANGEL_STATUS_PARAMETROS_RESULTADOS.xlsx'
REPORT_JSON_PATH = BASE_JSON_DIR / '09_ROADMAP_STATUS_REPORT_LATEST.json'

JSONS = [
    '00_RUN_ALL_REPORT_LATEST.json', '99_AI_CONTEXT_INDEX_LATEST.json', '00_RUN_STATE_LATEST.json',
    '00_01_MACHINE_PROFILE_LATEST.json', '00_02_PYTHON_ENVIRONMENT_LATEST.json', '00_03_BASE_ARQUIVOS_LATEST.json',
    '01_MAPA_ATIVOS_LATEST.json', '01_CATALOGO_SERIES_LATEST.json',
    '02_01_DATA_QUALITY_REPORT_LATEST.json', '02_01_DATA_QUALITY_ROOT_CAUSE_LATEST.json',
    '03_FEATURES_CATALOG_LATEST.json', '03_FEATURES_RUN_REPORT_LATEST.json', '03_01_FEATURES_CUDA_BENCHMARK_LATEST.json', '03_FEATURES_RETRY_PLAN_LATEST.json',
    '04_LABELS_CATALOG_LATEST.json', '04_LABELS_RUN_REPORT_LATEST.json',
    '05_DATASETS_ML_LATEST.json',
    '06_WALK_FORWARD_LATEST.json',
    '07_BACKTEST_PORTFOLIO_LATEST.json', '07_BACKTEST_PARAM_SEARCH_LATEST.json', '07_BACKTEST_VALIDATION_LATEST.json', '07_BACKTEST_STRESS_LATEST.json', '07_BACKTEST_EXPERIMENT_REGISTRY_LATEST.json',
    '08_EXPERIMENT_REGISTRY_LATEST.json', '08_EXPERIMENT_REGISTRY_VALIDATION_LATEST.json',
    '09_ROADMAP_STATUS_REPORT_LATEST.json', '00_COST_MODEL.json'
]

def load_json(name: str) -> dict[str, Any]:
    path = BASE_JSON_DIR / name
    if not path.is_file():
        return {}
    last = None
    for _ in range(5):
        try:
            data = json.loads(path.read_text(encoding='utf-8'))
            return data if isinstance(data, dict) else {}
        except Exception as exc:
            last = exc
    return {'_read_error': str(last)}


def nested(data: dict[str, Any], *keys: str, default: Any = None) -> Any:
    cur: Any = data
    for key in keys:
        if not isinstance(cur, dict):
            return default
        cur = cur.get(key)
    return default if cur is None else cur


def fnum(value: Any, default: float | None = None) -> float | None:
    try:
        out = float(value)
    except Exception:
        return default
    return out if out == out else default


def fint(value: Any) -> str:
    try:
        return f'{int(value):,}'.replace(',', '.')
    except Exception:
        return 'n/d'


def fpct(value: Any) -> str:
    val = fnum(value)
    return 'n/d' if val is None else f'{val * 100:.2f}%'


def fdec(value: Any) -> str:
    val = fnum(value)
    return 'n/d' if val is None else f'{val:.4f}'


def inventory() -> list[dict[str, Any]]:
    rows = []
    for name in JSONS:
        path = BASE_JSON_DIR / name
        data = load_json(name)
        system = data.get('system', {}) if isinstance(data.get('system'), dict) else {}
        rows.append({
            'file': name,
            'exists': path.is_file(),
            'size_bytes': path.stat().st_size if path.is_file() else None,
            'mtime_local': datetime.fromtimestamp(path.stat().st_mtime).isoformat(timespec='seconds') if path.is_file() else None,
            'schema_version': data.get('schema_version'),
            'run_id': data.get('run_id') or system.get('run_id'),
            'generated_at': data.get('generated_at') or data.get('generated_at_utc') or system.get('generated_at') or system.get('generated_at_utc'),
            'read_error': data.get('_read_error'),
        })
    return rows


def mtime(name: str) -> float:
    path = BASE_JSON_DIR / name
    return path.stat().st_mtime if path.is_file() else 0.0


def compact(value: Any, limit: int = 500) -> Any:
    if value is None or isinstance(value, (int, float, bool)):
        return value
    if isinstance(value, (dict, list)):
        value = json.dumps(value, ensure_ascii=False, default=str)
    out = str(value)
    return out if len(out) <= limit else out[: limit - 15] + ' ...[truncado]'


MODULE_PY = [
    ('00', '00_RUN_ALL.py', 'Orquestra pipeline completo e atualiza status final'),
    ('00.01', '00_01_DIAGNOSTICO_HARDWARE_BARRAMENTO.py', 'Perfil de hardware, CUDA, barramento e Python'),
    ('00.03', '00_03_MAPEAMENTO_DIRETORIOS.py', 'Inventario de arquivos, pastas e manifestos'),
    ('02.01', '02_01_AUDITA_QUALIDADE_DADOS.py', 'Auditoria de qualidade, gaps e bloqueios ML'),
    ('1', '01_MAPA_ATIVOS.py', 'Mapa dinamico de ativos, series, periodicidade e arquivos'),
    ('2', '02_BUSCA_DADOS.py', 'Coleta e atualizacao de dados de mercado'),
    ('3', '03_GERA_FEATURES.py', 'Geracao de features e telemetria de performance'),
    ('4', '04_GERA_LABELS.py', 'Geracao de labels e auditoria anti-leakage'),
    ('5', '05_MONTA_DATASETS_ML.py', 'Join features/labels e montagem de datasets ML'),
    ('6', '06_WALK_FORWARD_TRAINING.py', 'Treino walk-forward com purging/embargo'),
    ('7', '07_BACKTEST_PORTFOLIO.py', 'Backtest portfolio realista com custos/stress/execucao'),
    ('8', '08_EXPERIMENT_REGISTRY.py', 'Registry formal de experimentos e linhagem'),
    ('09', '09_ATUALIZA_ROADMAP_STATUS.py', 'Comunicacao humana: Word, Excel e JSON de roadmap'),
]


def _first(row: dict[str, Any], *keys: str) -> Any:
    for key in keys:
        val = row.get(key)
        if val not in (None, ''):
            return val
    return None


def _date_from_range(value: Any, *keys: str) -> Any:
    if not isinstance(value, dict):
        return None
    for key in keys:
        val = value.get(key) or value.get(key.replace('_', '')) or value.get(key.upper())
        if val not in (None, ''):
            return val
    return None


def _module_jsons(script: str) -> list[str]:
    return {
        '00_RUN_ALL.py': ['00_RUN_ALL_REPORT_LATEST.json', '00_RUN_STATE_LATEST.json', '99_AI_CONTEXT_INDEX_LATEST.json'],
        '00_01_DIAGNOSTICO_HARDWARE_BARRAMENTO.py': ['00_01_MACHINE_PROFILE_LATEST.json', '00_02_PYTHON_ENVIRONMENT_LATEST.json'],
        '00_03_MAPEAMENTO_DIRETORIOS.py': ['00_03_BASE_ARQUIVOS_LATEST.json'],
        '02_01_AUDITA_QUALIDADE_DADOS.py': ['02_01_DATA_QUALITY_REPORT_LATEST.json', '02_01_DATA_QUALITY_ROOT_CAUSE_LATEST.json'],
        '01_MAPA_ATIVOS.py': ['01_MAPA_ATIVOS_LATEST.json', '01_CATALOGO_SERIES_LATEST.json'],
        '02_BUSCA_DADOS.py': ['01_MAPA_ATIVOS_LATEST.json', '00_03_BASE_ARQUIVOS_LATEST.json'],
        '03_GERA_FEATURES.py': ['03_FEATURES_CATALOG_LATEST.json', '03_FEATURES_RUN_REPORT_LATEST.json', '03_01_FEATURES_CUDA_BENCHMARK_LATEST.json'],
        '04_GERA_LABELS.py': ['04_LABELS_CATALOG_LATEST.json', '04_LABELS_RUN_REPORT_LATEST.json'],
        '05_MONTA_DATASETS_ML.py': ['05_DATASETS_ML_LATEST.json'],
        '06_WALK_FORWARD_TRAINING.py': ['06_WALK_FORWARD_LATEST.json'],
        '07_BACKTEST_PORTFOLIO.py': ['07_BACKTEST_PORTFOLIO_LATEST.json', '07_BACKTEST_STRESS_LATEST.json', '07_BACKTEST_VALIDATION_LATEST.json'],
        '08_EXPERIMENT_REGISTRY.py': ['08_EXPERIMENT_REGISTRY_LATEST.json'],
        '09_ATUALIZA_ROADMAP_STATUS.py': ['09_ROADMAP_STATUS_REPORT_LATEST.json'],
    }.get(script, [])


def _module_note(script: str, data: dict[str, dict[str, Any]]) -> str:
    if script == '01_MAPA_ATIVOS.py':
        s = data['01_MAPA_ATIVOS_LATEST.json'].get('summary', {})
        return f"{fint(s.get('total_series_parquet_mapped'))} series; {fint(s.get('total_ohlcv_series'))} OHLCV; {fint(s.get('warning_series_count'))} com warning."
    if script == '02_BUSCA_DADOS.py':
        return 'Sem run_report proprio dedicado; status inferido por inventario, mapa de ativos e arquivos parquet.'
    if script == '03_GERA_FEATURES.py':
        s = data['03_FEATURES_CATALOG_LATEST.json'].get('summary', {})
        backend = nested(data['03_FEATURES_RUN_REPORT_LATEST.json'], 'summary', 'feature_compute_backend')
        return f"{fint(s.get('series_ok'))} series OK; {fint(s.get('feature_catalog_count'))} features; backend {compact(backend, 160)}."
    if script == '04_GERA_LABELS.py':
        s = data['04_LABELS_CATALOG_LATEST.json'].get('summary', {})
        return f"{fint(s.get('series_ok'))} series OK; {fint(s.get('labels_catalog_count'))} labels; audit_pass {fint(s.get('audit_pass'))}."
    if script == '05_MONTA_DATASETS_ML.py':
        s = data['05_DATASETS_ML_LATEST.json'].get('summary', {})
        return f"{fint(s.get('datasets_ok'))} datasets; {fint(s.get('total_trainable_rows'))} linhas treinaveis; gate {compact(s.get('ml_gate_counts'), 160)}."
    if script == '06_WALK_FORWARD_TRAINING.py':
        s = data['06_WALK_FORWARD_LATEST.json'].get('summary', {})
        return f"{fint(s.get('experiments_ok'))} experimentos; balanced accuracy {fdec(s.get('avg_balanced_accuracy'))}; CUDA {s.get('cuda_enabled_for_this_run')}."
    if script == '07_BACKTEST_PORTFOLIO.py':
        s = data['07_BACKTEST_PORTFOLIO_LATEST.json'].get('summary', {})
        return f"{fint(s.get('portfolio_total_trades'))} trades; CAGR {fpct(s.get('portfolio_cagr'))}; retorno {fpct(s.get('portfolio_total_return'))}; DD {fpct(s.get('portfolio_max_drawdown'))}."
    if script == '08_EXPERIMENT_REGISTRY.py':
        s = data['08_EXPERIMENT_REGISTRY_LATEST.json'].get('summary', {})
        return f"{fint(s.get('registry_rows'))} linhas; validacao {s.get('validation_status')}; missing artifacts {s.get('missing_artifacts')}."
    if script == '00_01_DIAGNOSTICO_HARDWARE_BARRAMENTO.py':
        s = data['00_02_PYTHON_ENVIRONMENT_LATEST.json'].get('summary', {})
        return f"PyTorch CUDA={s.get('cuda_ready_for_pytorch')}; CuPy CUDA={s.get('cuda_ready_for_cupy')}."
    if script == '00_RUN_ALL.py':
        s = data['00_RUN_ALL_REPORT_LATEST.json'].get('summary', {})
        return f"Status {s.get('status')}; etapas OK {fint(s.get('stages_ok'))}; erros {fint(s.get('stages_error'))}."
    return 'Operacional dentro do pipeline; detalhes consolidados nas abas especificas.'


def _module_change(script: str) -> str:
    return {
        '01_MAPA_ATIVOS.py': 'Adicionar classificacao MVP/experimental, mercado/exchange e regras de elegibilidade por serie.',
        '02_BUSCA_DADOS.py': 'Criar JSON proprio de run: fontes chamadas, latencia, candles novos, retries, falhas por exchange e cobertura por ativo.',
        '03_GERA_FEATURES.py': 'Migrar blocos rolling caros para CUDA/CuPy com benchmark antes/depois e equivalencia numerica.',
        '04_GERA_LABELS.py': 'Adicionar labels por regime/custo/funding e validacao mais dura de anti-leakage.',
        '05_MONTA_DATASETS_ML.py': 'Separar datasets por familia de features e criar datasets de ablation controlados.',
        '06_WALK_FORWARD_TRAINING.py': 'Registrar importancia/estabilidade por janela e degradacao por custos.',
        '07_BACKTEST_PORTFOLIO.py': 'Calibrar microestrutura com dados reais de liquidez quando disponiveis.',
        '08_EXPERIMENT_REGISTRY.py': 'Adicionar decisao humana: pesquisar, repetir, rejeitar, arquivar, candidato testnet.',
        '00_RUN_ALL.py': 'Manter hook final de status e nunca misturar live trading com pesquisa.',
    }.get(script, 'Documentar parametros editaveis e impacto esperado.')


def build_module_details(data: dict[str, dict[str, Any]]) -> dict[str, list[dict[str, Any]]]:
    mapa = data['01_MAPA_ATIVOS_LATEST.json']; dq = data['02_01_DATA_QUALITY_REPORT_LATEST.json']; feat = data['03_FEATURES_CATALOG_LATEST.json']
    labels = data['04_LABELS_CATALOG_LATEST.json']; ds = data['05_DATASETS_ML_LATEST.json']; wf = data['06_WALK_FORWARD_LATEST.json']
    bt = data['07_BACKTEST_PORTFOLIO_LATEST.json']; stress = data['07_BACKTEST_STRESS_LATEST.json']; reg = data['08_EXPERIMENT_REGISTRY_LATEST.json']
    run = data['00_RUN_ALL_REPORT_LATEST.json']
    run_stages = {str(x.get('script') or x.get('stage') or x.get('name')): x for x in run.get('stages', []) if isinstance(x, dict)}
    modules = []
    for code, script, purpose in MODULE_PY:
        p = RULES_DIR / script; st = run_stages.get(script, {})
        modules.append({'module': code, 'python_file': script, 'exists': p.is_file(), 'last_modified': datetime.fromtimestamp(p.stat().st_mtime).isoformat(timespec='seconds') if p.is_file() else None, 'size_kb': round(p.stat().st_size/1024, 1) if p.is_file() else None, 'purpose': purpose, 'last_run_status': st.get('status') or st.get('result'), 'elapsed_seconds': st.get('elapsed_seconds') or st.get('duration_seconds'), 'json_outputs': ', '.join(_module_jsons(script)), 'current_assessment': _module_note(script, data), 'what_can_change': _module_change(script)})

    q_by_id = dq.get('series', {}) if isinstance(dq.get('series'), dict) else {}
    assets_series = []
    fetch = {}
    for s in mapa.get('series_catalog', []) if isinstance(mapa.get('series_catalog'), list) else []:
        q = q_by_id.get(s.get('series_id'), {}) if isinstance(q_by_id, dict) else {}; dr = s.get('date_range') if isinstance(s.get('date_range'), dict) else {}; qual = s.get('quality') if isinstance(s.get('quality'), dict) else {}
        rows_val = _first(s, 'rows', 'row_count') or qual.get('rows') or q.get('rows')
        assets_series.append({'series_id': s.get('series_id'), 'asset': s.get('asset'), 'symbol': s.get('symbol'), 'source': s.get('source'), 'dataset_kind': s.get('dataset_kind'), 'timeframe': s.get('timeframe'), 'periodicity': compact(s.get('periodicity'), 220), 'date_start': _date_from_range(dr, 'min', 'start', 'first_datetime'), 'date_end': _date_from_range(dr, 'max', 'end', 'last_datetime'), 'rows': rows_val, 'quality_status': q.get('status') or qual.get('status') or s.get('status'), 'ml_quality_status': q.get('ml_quality_status') or q.get('ml_status') or qual.get('ml_quality_status'), 'root_cause': compact(q.get('root_cause') or q.get('root_causes') or q.get('fail_root_cause'), 220), 'warnings': compact(q.get('warnings') or q.get('warning_codes') or qual.get('warnings'), 260), 'file': s.get('file')})
        key = (s.get('source'), s.get('asset'), s.get('symbol'), s.get('dataset_kind'))
        item = fetch.setdefault(key, {'source': key[0], 'asset': key[1], 'symbol': key[2], 'dataset_kind': key[3], 'series_count': 0, 'timeframes': set(), 'rows_total': 0, 'date_start': None, 'date_end': None, 'status_note': 'Inferido do mapa de ativos; criar JSON proprio para coleta.'})
        item['series_count'] += 1
        if s.get('timeframe'): item['timeframes'].add(str(s.get('timeframe')))
        if isinstance(rows_val, (int, float)): item['rows_total'] += int(rows_val)
        start = _date_from_range(dr, 'min', 'start'); end = _date_from_range(dr, 'max', 'end')
        if start and (not item['date_start'] or str(start) < str(item['date_start'])): item['date_start'] = start
        if end and (not item['date_end'] or str(end) > str(item['date_end'])): item['date_end'] = end
    fetch_rows = []
    for item in fetch.values():
        item['timeframes'] = ', '.join(sorted(item['timeframes'])); fetch_rows.append(item)
    features_catalog = [{'feature': x.get('feature'), 'family': x.get('family'), 'type_feature': x.get('type_feature'), 'lookback': x.get('lookback'), 'uses_future_data': x.get('uses_future_data'), 'risk_relevance': x.get('risk_relevance'), 'ml_relevance': x.get('ml_relevance'), 'description': x.get('description'), 'formula': compact(x.get('formula'), 320)} for x in feat.get('feature_catalog', []) if isinstance(x, dict)]
    existing_families = Counter([str(x.get('family') or 'unknown') for x in feat.get('feature_catalog', []) if isinstance(x, dict)])
    feature_ideas = []
    for fam, examples, priority, care in [
        ('microstructure', 'order book imbalance, spread proxy, trade sign imbalance', 'Alta', 'Exige dados de book/trades por exchange; nao inventar sem fonte.'),
        ('funding_basis', 'funding momentum, funding z-score, basis perp/spot', 'Alta', 'Relevante para perp; precisa funding historico por exchange.'),
        ('cross_asset', 'BTC/ETH beta dinamico, lead-lag, dispersao setorial', 'Alta', 'Criar com cuidado para evitar vazamento temporal.'),
        ('regime', 'volatility state, trend/chop regime, liquidity drought', 'Alta', 'Usar para filtros e sizing, nao so como feature bruta.'),
        ('calendar', 'hora UTC, dia semana, sessao Asia/EUA, vencimentos futuros', 'Media', 'Util para cripto e commodities; validar por ablation.'),
        ('tail_risk', 'drawdown speed, downside semivariance, gap risk', 'Alta', 'Importante para limitar drawdown e stops.'),
        ('commodities_curve', 'front/second future spread, roll yield, contango/backwardation', 'Futura', 'Somente quando pipeline de futuros estiver ativo.'),
        ('sentiment_flow', 'OI delta, volume shock, liquidation proxy', 'Media', 'Depende de fonte confiavel e timestampavel.'),
    ]:
        feature_ideas.append({'family_candidate': fam, 'examples': examples, 'priority': priority, 'status_current': 'existente' if fam in existing_families else 'pendente/nao implementada', 'care': care, 'existing_feature_count_same_family': existing_families.get(fam, 0)})

    def _series_outputs(items, kind):
        rows = []
        for x in items if isinstance(items, list) else []:
            row = {'status': x.get('status'), 'asset': x.get('asset'), 'symbol': x.get('symbol'), 'source': x.get('source'), 'timeframe': x.get('timeframe'), 'series_id': x.get('series_id'), 'output_rows': x.get('output_rows') or x.get('rows'), 'output_columns': x.get('output_columns') or x.get('columns'), 'elapsed_seconds': x.get('elapsed_seconds'), 'memory_mb_start': x.get('memory_mb_start'), 'memory_mb_end': x.get('memory_mb_end'), 'output_path': x.get('output_path')}
            if kind == 'features': row.update({'feature_columns_count': x.get('feature_columns_count'), 'backend': compact(x.get('feature_execution_profile') or x.get('compute_backend'), 180), 'quality_status': x.get('quality_status_normalized') or x.get('quality_ml_status')})
            if kind == 'labels': row.update({'label_columns_count': x.get('label_columns_count'), 'audit_status': x.get('audit_status')})
            rows.append(row)
        return rows

    labels_config = []
    for x in labels.get('label_configs', []) if isinstance(labels.get('label_configs'), list) else []:
        labels_config.append({'label_config_id': x.get('label_config_id'), 'label_type': x.get('label_type'), 'enabled': x.get('enabled'), 'horizons': compact(x.get('horizons'), 160), 'description': x.get('description'), 'side': None, 'uses_future_data': None, 'change_idea': 'Testar ablation por horizonte, custo e regime antes de ampliar.'})
    for x in labels.get('labels_catalog', []) if isinstance(labels.get('labels_catalog'), list) else []:
        labels_config.append({'label_config_id': x.get('label_config_id'), 'label_type': x.get('label_type'), 'enabled': True, 'horizons': x.get('horizon_bars'), 'description': x.get('label_name'), 'side': x.get('side'), 'uses_future_data': x.get('uses_future_data'), 'change_idea': 'Catalogado; avaliar contribuicao via dataset/walk-forward/backtest.'})

    datasets = []
    for x in ds.get('datasets', []) if isinstance(ds.get('datasets'), list) else []:
        datasets.append({'status': x.get('status'), 'asset': x.get('asset'), 'symbol': x.get('symbol'), 'source': x.get('source'), 'timeframe': x.get('timeframe'), 'rows': x.get('rows'), 'trainable_rows': x.get('trainable_rows') or x.get('ml_trainable_rows'), 'columns': x.get('columns'), 'feature_columns_count': x.get('feature_columns_count'), 'xasset_columns_count': x.get('xasset_columns_count'), 'label_columns_count': x.get('label_columns_count'), 'allowed_feature_columns_count': x.get('allowed_feature_columns_count'), 'ml_gate': x.get('ml_gate') or x.get('ml_quality_gate'), 'quality_status': x.get('quality_status') or x.get('dataset_quality_status'), 'root_cause': compact(x.get('root_cause') or x.get('ml_root_cause'), 180), 'output_path': x.get('output_path')})

    wf_rows = []
    for x in wf.get('experiments', []) if isinstance(wf.get('experiments'), list) else []:
        m = x.get('metrics') if isinstance(x.get('metrics'), dict) else {}
        wf_rows.append({'experiment_id': x.get('experiment_id'), 'status': x.get('status'), 'asset': x.get('asset'), 'symbol': x.get('symbol'), 'source': x.get('source'), 'timeframe': x.get('timeframe'), 'target_col': x.get('target_col'), 'model_type': x.get('model_type'), 'backend': compact(x.get('backend_plan'), 180), 'rows_loaded': x.get('rows_loaded'), 'feature_columns_count': x.get('feature_columns_count'), 'windows_ok': x.get('windows_ok') or m.get('windows_ok'), 'balanced_accuracy': x.get('balanced_accuracy') or m.get('balanced_accuracy'), 'auc': x.get('auc') or m.get('auc'), 'oos_rows': x.get('oos_rows') or m.get('oos_rows'), 'artifact_path': x.get('experiment_dir') or x.get('predictions_path')})

    bt_rows = []
    for x in bt.get('experiments', []) if isinstance(bt.get('experiments'), list) else []:
        m = x.get('metrics') if isinstance(x.get('metrics'), dict) else {}
        bt_rows.append({'experiment_id': x.get('experiment_id'), 'status': x.get('status'), 'asset': x.get('asset'), 'symbol': x.get('symbol'), 'source': x.get('source'), 'timeframe': x.get('timeframe'), 'target_col': x.get('target_col'), 'horizon_bars': x.get('horizon_bars'), 'rows': x.get('rows'), 'trades': x.get('trades') or m.get('trades') or x.get('total_trades'), 'total_return': x.get('total_return') or m.get('total_return'), 'cagr': x.get('cagr') or m.get('cagr'), 'max_drawdown': x.get('max_drawdown') or m.get('max_drawdown'), 'win_rate': x.get('win_rate') or m.get('win_rate'), 'profit_factor': x.get('profit_factor') or m.get('profit_factor'), 'risk_controls': compact(x.get('risk_controls'), 240), 'side_distribution': compact(x.get('side_distribution'), 160), 'trades_path': x.get('trades_path'), 'equity_path': x.get('equity_path')})

    scenarios = stress.get('scenarios') or nested(bt, 'stress_tests', 'scenarios') or []
    stress_rows = [{'scenario': x.get('scenario') or x.get('name'), 'status': x.get('status'), 'total_return': x.get('total_return'), 'cagr': x.get('cagr'), 'max_drawdown': x.get('max_drawdown'), 'trades': x.get('trades'), 'notes': compact(x, 500)} for x in scenarios if isinstance(x, dict)]

    registry = []
    for x in reg.get('experiments', []) if isinstance(reg.get('experiments'), list) else []:
        registry.append({'registry_id': x.get('registry_id'), 'experiment_id': x.get('experiment_id'), 'status': x.get('status') or x.get('experiment_status'), 'asset': x.get('asset'), 'symbol': x.get('symbol'), 'source': x.get('source'), 'timeframe': x.get('timeframe'), 'target_col': x.get('target_col'), 'dataset_trainable_rows': x.get('dataset_trainable_rows'), 'dataset_ml_gate': x.get('dataset_ml_gate'), 'wf_balanced_accuracy': x.get('wf_balanced_accuracy') or x.get('walk_forward_balanced_accuracy'), 'backtest_total_return': x.get('backtest_total_return'), 'backtest_cagr': x.get('backtest_cagr'), 'backtest_max_drawdown': x.get('backtest_max_drawdown'), 'approval_status': x.get('approval_status'), 'portfolio_research_status': x.get('portfolio_research_status'), 'config_hash': x.get('config_hash'), 'dataset_fingerprint': x.get('dataset_fingerprint'), 'model_fingerprint': x.get('model_fingerprint')})

    parameters = []
    for source, obj, note in [('3_FEATURES', feat.get('feature_config', {}), 'Feature config atual'), ('3_PERFORMANCE', feat.get('performance_policy', {}), 'Hardware/CPU/CUDA features'), ('4_LABEL_CONFIG', labels.get('label_configs', []), 'Labels e horizontes'), ('5_POLICY', ds.get('policy', {}), 'Join, allowed features, anti-leakage'), ('6_CONFIG', wf.get('config', {}), 'Treino walk-forward'), ('6_BACKEND', wf.get('backend_plan', {}), 'CUDA/CPU training'), ('7_CONFIG', bt.get('config', {}), 'Backtest portfolio'), ('7_POLICY', bt.get('policy', {}), 'Execucao, custos e live guard'), ('8_POLICY', reg.get('policy', {}), 'Registry e aprovacao'), ('COST_MODEL', data['00_COST_MODEL.json'], 'Custos/funding/slippage assumptions')]:
        if isinstance(obj, dict):
            for k, v in obj.items(): parameters.append({'source': source, 'parameter': k, 'value': compact(v, 700), 'meaning': note, 'can_change': _parameter_change_note(source, k)})
        elif isinstance(obj, list):
            for i, v in enumerate(obj): parameters.append({'source': source, 'parameter': f'item_{i+1}', 'value': compact(v, 700), 'meaning': note, 'can_change': 'Alterar somente com ablation e registro no registry.'})

    improvements = [
        {'priority': 1, 'module': '02_BUSCA_DADOS.py', 'improvement': 'Criar run_report dedicado para coleta', 'why': 'Hoje a coleta e inferida por arquivos/series; falta telemetria de exchange, candles novos, retries e latencia.', 'risk_if_ignored': 'Codex e usuario nao conseguem auditar claramente falhas de dados.', 'status': 'pendente'},
        {'priority': 2, 'module': '03_GERA_FEATURES.py', 'improvement': 'CUDA em blocos rolling caros', 'why': 'Ultimo run de features ainda mostra pandas_cpu; ha volume alto de matriz.', 'risk_if_ignored': 'Tempo alto limita iteracao e ablation.', 'status': 'proximo'},
        {'priority': 3, 'module': 'Features', 'improvement': 'Ablation por familias e features nao convencionais', 'why': 'Backtest atual e negativo; precisamos achar ou rejeitar edge com disciplina.', 'risk_if_ignored': 'Aumentar complexidade sem saber o que ajuda.', 'status': 'proximo'},
        {'priority': 4, 'module': '07_BACKTEST_PORTFOLIO.py', 'improvement': 'Calibrar microestrutura com dados reais', 'why': 'Slippage/funding assumidos sao bons para stress, mas devem convergir para dados por exchange.', 'risk_if_ignored': 'Backtest pode ficar distante da execucao real.', 'status': 'pendente'},
        {'priority': 5, 'module': '08_EXPERIMENT_REGISTRY.py', 'improvement': 'Adicionar campo decisao humana', 'why': 'Separar pesquisa tecnica, candidato promissor, rejeitado, arquivado e pronto para testnet.', 'risk_if_ignored': 'Historico vira deposito, nao ferramenta de decisao.', 'status': 'pendente'},
        {'priority': 6, 'module': 'Execucao testnet', 'improvement': 'Conectar Bybit/Binance/Kraken Pro em ambiente de teste', 'why': 'Validar latencia, ordens, partial fill e robustez operacional.', 'risk_if_ignored': 'Modelo pode parecer bom mas falhar na execucao.', 'status': 'depois'},
    ]
    return {'modules': modules, 'assets_series': assets_series, 'data_fetch_status': sorted(fetch_rows, key=lambda x: (str(x.get('source')), str(x.get('asset')), str(x.get('dataset_kind')))), 'features_catalog': features_catalog, 'features_series': _series_outputs(feat.get('series_outputs', []), 'features'), 'features_ideas': feature_ideas, 'labels_config': labels_config, 'labels_series': _series_outputs(labels.get('series_outputs', []), 'labels'), 'datasets': datasets, 'walk_forward': wf_rows, 'backtest_experiments': bt_rows, 'stress_tests': stress_rows, 'registry': registry, 'parameters': parameters, 'improvements': improvements}


def _parameter_change_note(source: str, key: str) -> str:
    k = str(key).lower()
    if any(x in k for x in ['threshold', 'stop', 'take', 'risk', 'leverage', 'drawdown', 'position', 'exposure']): return 'Editavel para pesquisa; precisa backtest/stress e config_hash.'
    if any(x in k for x in ['cuda', 'worker', 'memory', 'parallel']): return 'Editavel para performance; validar equivalencia numerica e tempo antes/depois.'
    if any(x in k for x in ['target', 'label', 'horizon']): return 'Editavel, mas altera problema de ML; exige novo dataset/walk-forward/backtest.'
    if source == 'COST_MODEL': return 'Editavel como stress/assumption; nunca reduzir custo para fabricar resultado.'
    return 'Mudanca possivel, mas registrar motivo e impacto esperado.'


def assess() -> dict[str, Any]:
    data = {name: load_json(name) for name in JSONS}
    inv = inventory()
    run = data['00_RUN_ALL_REPORT_LATEST.json']
    mapa = data['01_MAPA_ATIVOS_LATEST.json']; dq = data['02_01_DATA_QUALITY_REPORT_LATEST.json']
    feat = data['03_FEATURES_CATALOG_LATEST.json']; feat_run = data['03_FEATURES_RUN_REPORT_LATEST.json']
    labels = data['04_LABELS_CATALOG_LATEST.json']; ds = data['05_DATASETS_ML_LATEST.json']; wf = data['06_WALK_FORWARD_LATEST.json']
    bt = data['07_BACKTEST_PORTFOLIO_LATEST.json']; val = data['07_BACKTEST_VALIDATION_LATEST.json']; stress = data['07_BACKTEST_STRESS_LATEST.json']
    reg = data['08_EXPERIMENT_REGISTRY_LATEST.json']; pyenv = data['00_02_PYTHON_ENVIRONMENT_LATEST.json']; bench = data['03_01_FEATURES_CUDA_BENCHMARK_LATEST.json']
    run_s = run.get('summary', {}); dq_s = dq.get('summary', {}); feat_s = feat.get('summary', {}); feat_run_s = feat_run.get('summary', {})
    label_s = labels.get('summary', {}); ds_s = ds.get('summary', {}); wf_s = wf.get('summary', {}); bt_s = bt.get('summary', {}); reg_s = reg.get('summary', {})
    run_time = mtime('00_RUN_ALL_REPORT_LATEST.json'); feat_time = max(mtime('03_FEATURES_CATALOG_LATEST.json'), mtime('03_FEATURES_RUN_REPORT_LATEST.json'))
    downstream = {'labels': mtime('04_LABELS_CATALOG_LATEST.json'), 'datasets': mtime('05_DATASETS_ML_LATEST.json'), 'walk_forward': mtime('06_WALK_FORWARD_LATEST.json'), 'backtest': mtime('07_BACKTEST_PORTFOLIO_LATEST.json'), 'registry': mtime('08_EXPERIMENT_REGISTRY_LATEST.json')}
    stale = [k for k, v in downstream.items() if feat_time and v and v + 30 < feat_time]
    self_generated = {'09_ROADMAP_STATUS_REPORT_LATEST.json', '99_AI_CONTEXT_INDEX_LATEST.json'}
    newer = [r['file'] for r in inv if r['file'] not in self_generated and (BASE_JSON_DIR / r['file']).is_file() and (BASE_JSON_DIR / r['file']).stat().st_mtime > run_time + 30]
    cagr = fnum(bt_s.get('portfolio_cagr')); total_ret = fnum(bt_s.get('portfolio_total_return')); max_dd = fnum(bt_s.get('portfolio_max_drawdown'))
    bal = fnum(wf_s.get('avg_balanced_accuracy'))
    red = []
    if newer: red.append('Há JSONs mais novos que o último relatório do 00_RUN_ALL; este snapshot pode ser parcial ou refletir run ainda em andamento.')
    if stale: red.append('A etapa 3/features está mais nova que downstreams; labels/datasets/walk-forward/backtest/registry podem estar defasados até o run atual terminar.')
    if cagr is not None and cagr < 0: red.append('Backtest de portfólio atual tem CAGR negativo; não existe evidência econômica para testnet por performance.')
    if max_dd is not None and abs(max_dd) > 0.08: red.append('Max drawdown observado ultrapassa a referência de 8%; isto bloqueia qualquer leitura otimista.')
    if bal is not None and bal < 0.55: red.append('Balanced accuracy média está apenas levemente acima de 50%; edge estatístico ainda é fraco.')
    if dq_s.get('ml_caution_series_count', 0): red.append('Muitas séries seguem em ML_CAUTION, principalmente por TIME_GAPS; isso contamina confiança nos resultados.')
    verdict = 'SNAPSHOT PARCIAL: AGUARDAR FIM DO 00_RUN_ALL ANTES DE DECIDIR' if newer else 'PESQUISA AVANÇOU, MAS NÃO HÁ SINAL ECONÔMICO APROVADO'
    stages = [
        ['Git/GitHub','OK','Repo local ligado ao origin/main e CI leve passou.','Repo é público; dados/JSONs/docx continuam fora.','Commits pequenos antes de mudanças grandes.'],
        ['Dados','PARCIAL','{} séries mapeadas; {} OHLCV.'.format(fint(nested(mapa,'summary','total_series_parquet_mapped')), fint(nested(mapa,'summary','total_ohlcv_series'))),'Warnings e séries unknown ainda existem.','Isolar universo MVP.'],
        ['Qualidade','USÁVEL COM CAUTELA','{} ML_READY, {} ML_CAUTION, {} ML_BLOCKED.'.format(fint(dq_s.get('ml_ready_series_count')), fint(dq_s.get('ml_caution_series_count')), fint(dq_s.get('ml_blocked_series_count'))),'TIME_GAPS é raiz dominante.','Priorizar séries realmente usadas.'],
        ['Features','OPERACIONAL','{} séries OK; {} features; backend atual {}.'.format(fint(feat_s.get('series_ok')), fint(feat_s.get('feature_catalog_count')), nested(feat_run_s,'feature_compute_backend','series_backend_counts')),'Último run de features aparece em CPU/pandas.','CUDA só com benchmark e fallback.'],
        ['Labels','OPERACIONAL/DEFASAGEM POSSÍVEL','{} séries; {} labels.'.format(fint(label_s.get('series_ok')), fint(label_s.get('labels_catalog_count'))),'Pode estar defasado se features acabou de atualizar.','Confirmar timestamps após o run.'],
        ['Datasets ML','OPERACIONAL, MAS CAUTION','{} datasets; {} linhas treináveis.'.format(fint(ds_s.get('datasets_ok')), fint(ds_s.get('total_trainable_rows'))),'Todos atuais são ML_CAUTION_ACCEPTABLE.','Medir robustez por causa raiz.'],
        ['Walk-forward','TÉCNICO OK','{} experimentos; balanced accuracy {}.'.format(fint(wf_s.get('experiments_ok')), fdec(wf_s.get('avg_balanced_accuracy'))),'Métrica fraca para edge operacional.','Usar como filtro, não aprovação.'],
        ['Backtest portfolio','MOTOR OK, RESULTADO RUIM','{} trades; CAGR {}; retorno {}; DD {}; validação {}.'.format(fint(bt_s.get('portfolio_total_trades')), fpct(cagr), fpct(total_ret), fpct(max_dd), nested(val,'summary','status')),'PnL líquido negativo e DD acima de 8%.','Usar para ablação/stress, não testnet.'],
        ['Stress','OK COMO TESTE','Pior cenário {}; retorno {}; DD {}.'.format(nested(stress,'summary','worst_scenario'), fpct(nested(stress,'summary','worst_total_return')), fpct(nested(stress,'summary','worst_max_drawdown'))),'Stress confirma fragilidade econômica atual.','Comparar novas features sob stress.'],
        ['Experiment registry','FORMALIZADO','{} experimentos; validação {}; status {}.'.format(fint(reg_s.get('registry_rows')), reg_s.get('validation_status'), reg_s.get('status')),'Registry registra evidência, não aprova execução.','Virar tabela mestre de pesquisa.'],
        ['CUDA','PRONTO CONTROLADO','PyTorch CUDA={}; CuPy CUDA={}; benchmark numérico OK={}.'.format(nested(pyenv,'summary','cuda_ready_for_pytorch'), nested(pyenv,'summary','cuda_ready_for_cupy'), nested(bench,'summary','numerical_comparison_ok')),'Acelerar sem equivalência numérica cria falso progresso.','Migrar apenas blocos caros.'],
    ]
    roadmap = [
        [1,'Manter DOCX/XLSX/JSON atualizados automaticamente ao fim de cada 00_RUN_ALL','A comunicação humana e IA precisa refletir o run final completo.','IMPLEMENTADO'],
        [2,'Criar run_report dedicado para 02_BUSCA_DADOS.py','A coleta ainda não é tão auditável quanto as etapas 3 a 8.','PRÓXIMO'],
        [3,'Migrar CUDA controlado em rolling windows pesadas da etapa 3','Reduz tempo de iteração sem alterar a tese econômica.','PRÓXIMO'],
        [4,'Criar fila de hipóteses e ablação formal de features','O backtest atual é negativo; precisamos saber o que ajuda ou atrapalha.','PRÓXIMO'],
        [5,'Ligar testnet Bybit/Binance/Kraken Pro','Validar execução, latência e partial fill sem risco real.','DEPOIS'],
        [6,'Adaptar para futuros de commodities: coffee, cocoa, cotton','Exige vencimento, rolagem, curva, calendário e liquidez.','FUTURO'],
    ]
    params = [
        ['target_annual_return_min', bt_s.get('target_annual_return_min'), '07_BACKTEST_PORTFOLIO_LATEST', 'Alvo econômico mínimo; não é aprovação automática.'],
        ['reference_drawdown_limit', bt_s.get('reference_drawdown_limit'), '07_BACKTEST_PORTFOLIO_LATEST', 'Referência temporária de 8%.'],
        ['approval_status', bt_s.get('approval_status'), '07_BACKTEST_PORTFOLIO_LATEST', 'Pesquisa separada de aprovação.'],
        ['feature_backend_mode', nested(feat_run_s,'feature_compute_backend','mode'), '03_FEATURES_RUN_REPORT_LATEST', 'Último run de features em auto.'],
        ['walk_forward_backend_usage', wf_s.get('backend_usage'), '06_WALK_FORWARD_LATEST', 'Treino já usa torch_cuda.'],
        ['cuda_ready_for_pytorch', nested(pyenv,'summary','cuda_ready_for_pytorch'), '00_02_PYTHON_ENVIRONMENT_LATEST', 'Pronto para uso controlado.'],
        ['cuda_ready_for_cupy', nested(pyenv,'summary','cuda_ready_for_cupy'), '00_02_PYTHON_ENVIRONMENT_LATEST', 'Pronto para blocos selecionados.'],
    ]
    return {'schema_version': SCHEMA_VERSION, 'system': {'name': 'ARCHANGEL', 'layer': 'ROADMAP_STATUS', 'script': SCRIPT_NAME, 'generated_at_utc': datetime.now(timezone.utc).isoformat(timespec='seconds')}, 'paths': {'docx_path': str(DOCX_PATH), 'xlsx_path': str(XLSX_PATH), 'report_json_path': str(REPORT_JSON_PATH), 'base_json_dir': str(BASE_JSON_DIR)}, 'summary': {'verdict': verdict, 'likely_run_in_progress_or_partial_snapshot': bool(newer), 'downstream_stale_after_features': stale, 'red_flags_count': len(red), 'pipeline_status_latest_completed': run_s.get('status'), 'features_ok': feat_s.get('series_ok'), 'labels_ok': label_s.get('series_ok'), 'datasets_ok': ds_s.get('datasets_ok'), 'walk_forward_experiments_ok': wf_s.get('experiments_ok'), 'walk_forward_avg_balanced_accuracy': wf_s.get('avg_balanced_accuracy'), 'portfolio_cagr': cagr, 'portfolio_total_return': total_ret, 'portfolio_max_drawdown': max_dd, 'formal_registry_status': reg_s.get('status'), 'formal_registry_rows': reg_s.get('registry_rows'), 'excel_detail_sheets_count': 20}, 'staleness': {'jsons_newer_than_last_run_all_report': newer, 'downstream_stale_after_features': stale}, 'red_flags': red, 'stage_assessment': [{'stage':r[0], 'status':r[1], 'evidence':r[2], 'risk':r[3], 'next_action':r[4]} for r in stages], 'roadmap': [{'priority':r[0], 'step':r[1], 'why':r[2], 'status':r[3]} for r in roadmap], 'parameters': [{'parameter':r[0], 'value':r[1], 'source':r[2], 'note':r[3]} for r in params], 'module_details': build_module_details(data), 'json_inventory': inv}


def write_json(path: Path, payload: dict[str, Any]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    tmp = path.with_suffix(path.suffix + '.tmp')
    tmp.write_text(json.dumps(payload, ensure_ascii=False, indent=2, default=str), encoding='utf-8')
    os.replace(tmp, path)


def set_font(run, bold=False, size=10.5, color='222222'):
    run.bold = bold; run.font.size = Pt(size); run.font.color.rgb = RGBColor.from_string(color); run.font.name = 'Calibri'
    run._element.get_or_add_rPr().rFonts.set(qn('w:ascii'), 'Calibri'); run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')


def shade(cell, fill):
    pr = cell._tc.get_or_add_tcPr(); shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), fill); pr.append(shd)


def table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER; table.autofit = False
    for row in table.rows:
        for i, w in enumerate(widths):
            if i < len(row.cells): row.cells[i].width = Inches(w); row.cells[i].vertical_alignment = WD_ALIGN_VERTICAL.TOP


def para(doc, text, bold_prefix=None, color='222222'):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(6); p.paragraph_format.line_spacing = 1.15
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix); set_font(r, True, 10.5, color); r = p.add_run(text[len(bold_prefix):]); set_font(r, False, 10.5, color)
    else: set_font(p.add_run(text), False, 10.5, color)


def bullet(doc, text, color='222222'):
    p = doc.add_paragraph(style='List Bullet'); p.paragraph_format.space_after = Pt(4); set_font(p.add_run(text), False, 10, color)


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level); p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(5)
    for r in p.runs: set_font(r, True, 16 if level == 1 else 12.5, '0B2545' if level == 1 else '1F4D78')


def matrix(doc, headers, rows, widths):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = 'Table Grid'
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h; shade(t.rows[0].cells[i], 'E8EEF5')
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row): cells[i].text = str(value)
    for row in t.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for r in p.runs: set_font(r, p is t.rows[0].cells[0].paragraphs[0], 8.8)
    table_width(t, widths)


def build_docx(report):
    doc = Document(); sec = doc.sections[0]; sec.top_margin = Inches(0.8); sec.bottom_margin = Inches(0.75); sec.left_margin = Inches(0.85); sec.right_margin = Inches(0.85)
    normal = doc.styles['Normal']; normal.font.name = 'Calibri'; normal.font.size = Pt(10.5)
    title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.LEFT; set_font(title.add_run('ARCHANGEL'), True, 24, '0B2545')
    sub = doc.add_paragraph(); set_font(sub.add_run('Roadmap interno v3 | Prompt reverso, status real e próximos passos'), False, 12, '555555')
    meta = doc.add_paragraph(); set_font(meta.add_run(f"Atualizado em {datetime.now().strftime('%d/%m/%Y %H:%M')} | Projeto local: {ROOT_DIR}"), False, 9, '555555')
    heading(doc, '1. Veredito Executivo')
    matrix(doc, ['Campo','Valor'], [['Veredito', report['summary']['verdict']], ['Snapshot', 'Parcial/em andamento' if report['summary']['likely_run_in_progress_or_partial_snapshot'] else 'Coerente com último run completo'], ['Meta econômica','Retorno líquido mínimo desejado acima de 20% ao ano; ainda não demonstrado.'], ['Comunicação','JSONs alimentam IA/Codex; Word e Excel alimentam decisão humana.'], ['Regra','Pesquisa não é aprovação de execução. Testnet só depois de candidato minimamente promissor.']], [1.8,4.7])
    heading(doc, '2. Alertas Que Não Devem Ser Ignorados')
    for r in report['red_flags']: bullet(doc, r, '9B1C1C')
    heading(doc, '3. Estado Atual Por Camada')
    matrix(doc, ['Camada','Estado','Evidência','Risco','Próxima ação'], [[x['stage'],x['status'],x['evidence'],x['risk'],x['next_action']] for x in report['stage_assessment']], [1.05,1.05,1.8,1.3,1.3])
    heading(doc, '4. Métricas Duras')
    s=report['summary']; matrix(doc, ['Métrica','Valor'], [['Features OK',fint(s.get('features_ok'))],['Labels OK',fint(s.get('labels_ok'))],['Datasets OK',fint(s.get('datasets_ok'))],['Walk-forward OK',fint(s.get('walk_forward_experiments_ok'))],['Balanced accuracy média',fdec(s.get('walk_forward_avg_balanced_accuracy'))],['Backtest CAGR',fpct(s.get('portfolio_cagr'))],['Backtest retorno total',fpct(s.get('portfolio_total_return'))],['Backtest max drawdown',fpct(s.get('portfolio_max_drawdown'))],['Registry formal',f"{s.get('formal_registry_status')} com {fint(s.get('formal_registry_rows'))} experimentos"]], [2.1,4.4])
    para(doc, 'Leitura cruelmente realista: a engenharia do pipeline está avançada, mas o resultado econômico atual é ruim. A próxima vitória é descobrir, via ablação e stress, se existe edge líquido robusto.', 'Leitura cruelmente realista:')
    heading(doc, '5. Roadmap Recomendado')
    for r in report['roadmap']: bullet(doc, f"{r['priority']}. {r['step']} [{r['status']}]. Motivo: {r['why']}")
    heading(doc, '6. Como Ler Este Status')
    para(doc, 'Leia este DOCX, depois 09_ROADMAP_STATUS_REPORT_LATEST.json, 99_AI_CONTEXT_INDEX_LATEST.json e o Excel de status. Peça sempre diagnóstico objetivo, preservando anti-leakage e separando pesquisa de aprovação.')
    heading(doc, '7. Excel Operacional')
    para(doc, 'O Excel agora é o painel humano detalhado do projeto. Ele possui abas por módulo Python, ativos/séries, coleta, features, labels, datasets, walk-forward, backtest, stress, registry, parâmetros e melhorias sugeridas.')
    heading(doc, '8. Coerência Dos JSONs')
    matrix(doc, ['Campo','Valor'], [['Downstreams defasados', ', '.join(report['staleness']['downstream_stale_after_features']) or 'nenhum'], ['JSONs mais novos que run_all', len(report['staleness']['jsons_newer_than_last_run_all_report'])]], [2.1,4.4])
    doc.save(DOCX_PATH)


def style_sheet(ws):
    ws.sheet_view.showGridLines = False; thin = Side(style='thin', color='DADCE0')
    for row in ws.iter_rows():
        for c in row:
            c.font = Font(name='Calibri', size=10, color='222222'); c.alignment = Alignment(vertical='top', wrap_text=True); c.border = Border(bottom=thin)
            if c.row == 1: c.font = Font(name='Calibri', size=10, bold=True, color='FFFFFF'); c.fill = PatternFill('solid', fgColor='0B2545')
    ws.freeze_panes = 'A2'; ws.auto_filter.ref = ws.dimensions
    if ws.max_row >= 2 and ws.max_column >= 1:
        try:
            headers = [ws.cell(1, c).value for c in range(1, ws.max_column + 1)]
            if all(isinstance(h, str) and h.strip() for h in headers):
                ref = f'A1:{get_column_letter(ws.max_column)}{ws.max_row}'
                tname = ''.join(ch for ch in ws.title if ch.isalnum())[:24] or 'Tabela'
                tab = Table(displayName=f'{tname}Tbl', ref=ref)
                tab.tableStyleInfo = TableStyleInfo(name='TableStyleMedium2', showFirstColumn=False, showLastColumn=False, showRowStripes=True, showColumnStripes=False)
                ws.add_table(tab)
        except Exception:
            pass
    for i in range(1, ws.max_column+1):
        width = 10
        for j in range(1, min(ws.max_row,250)+1):
            v = ws.cell(j,i).value
            if v is not None: width = max(width, min(60, len(str(v))+2))
        ws.column_dimensions[get_column_letter(i)].width = width



def rowsheet(wb, name, headers, rows):
    ws = wb.create_sheet(name[:31]); ws.append(headers)
    for r in rows:
        ws.append([compact(r.get(h), 32000) for h in headers])
    style_sheet(ws); return ws


def build_xlsx(report):
    wb = Workbook(); ws = wb.active; ws.title = 'Dashboard'; ws.append(['Campo','Valor','Leitura'])
    s = report['summary']; details = report.get('module_details') or build_module_details({name: load_json(name) for name in JSONS})
    dashboard_rows = [
        ('Veredito', s.get('verdict'), 'Decisão executiva do snapshot'),
        ('Pipeline status', s.get('pipeline_status_latest_completed'), 'Último 00_RUN_ALL'),
        ('Features OK', s.get('features_ok'), 'Volume operacional de features'),
        ('Labels OK', s.get('labels_ok'), 'Targets gerados'),
        ('Datasets OK', s.get('datasets_ok'), 'Bases treináveis'),
        ('Walk-forward experiments OK', s.get('walk_forward_experiments_ok'), 'Treinos concluídos'),
        ('Balanced accuracy média', s.get('walk_forward_avg_balanced_accuracy'), 'Edge estatístico ainda fraco se perto de 0.50'),
        ('Portfolio CAGR', s.get('portfolio_cagr'), 'Resultado líquido anualizado do motor atual'),
        ('Portfolio retorno total', s.get('portfolio_total_return'), 'Resultado líquido acumulado'),
        ('Portfolio max drawdown', s.get('portfolio_max_drawdown'), 'Risco realizado no backtest'),
        ('Registry rows', s.get('formal_registry_rows'), 'Experimentos versionados'),
        ('Abas detalhadas', 20, 'Cobertura operacional humana'),
    ]
    for row in dashboard_rows: ws.append(list(row))
    ws['E1']='Métrica'; ws['F1']='Valor'
    metrics=[('Balanced Accuracy',s.get('walk_forward_avg_balanced_accuracy')),('CAGR',s.get('portfolio_cagr')),('Total Return',s.get('portfolio_total_return')),('Max Drawdown',s.get('portfolio_max_drawdown'))]
    for i,(k,v) in enumerate(metrics,2): ws.cell(i,5).value=k; ws.cell(i,6).value=v
    chart=BarChart(); chart.title='Métricas principais'; chart.add_data(Reference(ws,min_col=6,min_row=1,max_row=5),titles_from_data=True); chart.set_categories(Reference(ws,min_col=5,min_row=2,max_row=5)); chart.height=7; chart.width=14; ws.add_chart(chart,'H2')
    style_sheet(ws)

    rowsheet(wb,'00_Modulos_Python',['module','python_file','exists','last_modified','size_kb','purpose','last_run_status','elapsed_seconds','json_outputs','current_assessment','what_can_change'],details['modules'])
    rowsheet(wb,'1_Mapa_Ativos_Series',['series_id','asset','symbol','source','dataset_kind','timeframe','periodicity','date_start','date_end','rows','quality_status','ml_quality_status','root_cause','warnings','file'],details['assets_series'])
    rowsheet(wb,'2_Busca_Dados_Status',['source','asset','symbol','dataset_kind','series_count','timeframes','rows_total','date_start','date_end','status_note'],details['data_fetch_status'])
    rowsheet(wb,'3_Features_Catalogo',['feature','family','type_feature','lookback','uses_future_data','risk_relevance','ml_relevance','description','formula'],details['features_catalog'])
    rowsheet(wb,'3_Features_Series',['status','asset','symbol','source','timeframe','series_id','output_rows','output_columns','feature_columns_count','backend','quality_status','elapsed_seconds','memory_mb_start','memory_mb_end','output_path'],details['features_series'])
    rowsheet(wb,'3_Features_Ideias',['family_candidate','examples','priority','status_current','care','existing_feature_count_same_family'],details['features_ideas'])
    rowsheet(wb,'4_Labels_Config',['label_config_id','label_type','enabled','horizons','description','side','uses_future_data','change_idea'],details['labels_config'])
    rowsheet(wb,'4_Labels_Series',['status','asset','symbol','source','timeframe','series_id','output_rows','output_columns','label_columns_count','audit_status','elapsed_seconds','memory_mb_start','memory_mb_end','output_path'],details['labels_series'])
    rowsheet(wb,'5_Datasets',['status','asset','symbol','source','timeframe','rows','trainable_rows','columns','feature_columns_count','xasset_columns_count','label_columns_count','allowed_feature_columns_count','ml_gate','quality_status','root_cause','output_path'],details['datasets'])
    rowsheet(wb,'6_Walk_Forward',['experiment_id','status','asset','symbol','source','timeframe','target_col','model_type','backend','rows_loaded','feature_columns_count','windows_ok','balanced_accuracy','auc','oos_rows','artifact_path'],details['walk_forward'])
    rowsheet(wb,'7_Backtest_Experimentos',['experiment_id','status','asset','symbol','source','timeframe','target_col','horizon_bars','rows','trades','total_return','cagr','max_drawdown','win_rate','profit_factor','risk_controls','side_distribution','trades_path','equity_path'],details['backtest_experiments'])
    rowsheet(wb,'7_Stress_Tests',['scenario','status','total_return','cagr','max_drawdown','trades','notes'],details['stress_tests'])
    rowsheet(wb,'8_Registry',['registry_id','experiment_id','status','asset','symbol','source','timeframe','target_col','dataset_trainable_rows','dataset_ml_gate','wf_balanced_accuracy','backtest_total_return','backtest_cagr','backtest_max_drawdown','approval_status','portfolio_research_status','config_hash','dataset_fingerprint','model_fingerprint'],details['registry'])
    rowsheet(wb,'Parametros_Editaveis',['source','parameter','value','meaning','can_change'],details['parameters'])
    rowsheet(wb,'Melhorias_Sugeridas',['priority','module','improvement','why','risk_if_ignored','status'],details['improvements'])
    rowsheet(wb,'Pipeline_Status',['stage','status','evidence','risk','next_action'],report['stage_assessment'])
    rowsheet(wb,'Roadmap',['priority','step','why','status'],report['roadmap'])
    rowsheet(wb,'Parametros_Resumo',['parameter','value','source','note'],report['parameters'])
    rowsheet(wb,'JSON_Inventory',['file','exists','size_bytes','mtime_local','schema_version','run_id','generated_at','read_error'],report['json_inventory'])
    flags = wb.create_sheet('Red_Flags'); flags.append(['red_flag']); [flags.append([x]) for x in report['red_flags']]; style_sheet(flags)
    for sheet in wb.worksheets:
        sheet.sheet_view.showGridLines = False
        if sheet.max_column >= 1: sheet.freeze_panes = 'A2'
        for row in sheet.iter_rows(min_row=2):
            for cell in row:
                if isinstance(cell.value, float):
                    header = str(sheet.cell(1, cell.column).value).lower()
                    cell.number_format = '0.00%' if any(k in header for k in ['return','cagr','drawdown','accuracy','auc','rate','pct']) else '0.0000'
                elif isinstance(cell.value, int): cell.number_format = '#,##0'
    wb.save(XLSX_PATH)


def update_ai_index(report):
    path = BASE_JSON_DIR/'99_AI_CONTEXT_INDEX_LATEST.json'
    if not path.is_file(): return
    idx = load_json('99_AI_CONTEXT_INDEX_LATEST.json')
    idx.setdefault('summary', {})['roadmap_status_verdict'] = report['summary']['verdict']
    idx['summary']['roadmap_status_generated_at_utc'] = report['system']['generated_at_utc']
    order = idx.setdefault('ai_reading_order', [])
    if '09_ROADMAP_STATUS_REPORT_LATEST.json' not in order: order.insert(1, '09_ROADMAP_STATUS_REPORT_LATEST.json')
    write_json(path, idx)


def main() -> int:
    BASE_JSON_DIR.mkdir(parents=True, exist_ok=True)
    report = assess(); write_json(REPORT_JSON_PATH, report); build_docx(report); build_xlsx(report); update_ai_index(report)
    print('ARCHANGEL roadmap/status atualizado')
    print(f'DOCX: {DOCX_PATH}')
    print(f'XLSX: {XLSX_PATH}')
    print(f'JSON: {REPORT_JSON_PATH}')
    print(f"Veredito: {report['summary']['verdict']}")
    return 0

if __name__ == '__main__':
    raise SystemExit(main())




